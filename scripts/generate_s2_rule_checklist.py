from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def _add_note(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(4)


def _add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        header_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in header_cells[index].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
    document.add_paragraph("")


def build_checklist(output_path: Path) -> None:
    document = Document()
    document.styles["Normal"].font.name = "Microsoft YaHei"
    document.styles["Normal"].font.size = Pt(10)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        document.styles[style_name].font.name = "Microsoft YaHei"

    title = document.add_heading("CableTrayAI S2 方钢托臂规则确认清单", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_note(document, f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    _add_note(
        document,
        "用途：给科室审核建模、计算、结果提取、评定和报告注入逻辑。正式结论仍以 ANSYS 输出、Excel/确定性公式和 source_ref 为准；大模型只做解释、质控和修复建议。",
    )
    _add_note(
        document,
        "来源：用户批注版 S2_RULE_CONFIRMATION_CHECKLIST.docx、新版“电缆桥架--.xlsx”、source_materials 中标准命令流/PIP/MAC/SECT、既有报告和真实 ANSYS 输出。",
    )

    _add_note(document, "0. 大白话术语表")
    _add_table(
        document,
        ["名称", "大白话解释", "主要用途", "出问题说明"],
        [
            ["MAXBEAMSTRESS.LIS", "支架/方钢梁单元应力结果文件", "表6-1 支架方钢应力评定：拉伸、压缩、弯曲、剪切、拉弯、压弯", "缺失、全零或集合不对时，支架方钢评定不能出正式结论"],
            ["TMAXBEAMSTRESS.LIS", "托臂梁单元应力结果文件", "托臂应力评定；小截面托臂根部焊缝等效应力表", "小截面任务全零时，表6-2-1 和附录C不能通过"],
            ["JCZH.LIS", "支架基础载荷文件", "表6-3/对应基础载荷：FX、FY、FZ、MY、MZ", "UNKNOWN 节点、全零或少工况时，基础载荷表不能用"],
            ["LS-FORCE.LIS", "托盘与托臂连接螺栓载荷文件", "连接螺栓载荷和螺栓应力评定", "不能用接近报告的节点硬凑；必须来自固定点集合"],
            ["HF-FORCE.LIS", "大截面托臂根部焊缝受力文件", "方钢外边长 >120mm 时输出表6-2 和表6-3", "小截面不强制依赖它，大截面缺失则阻断"],
            ["Mode.oup", "模态频率输出文件", "附录A 模态频率；校核 MT 是否覆盖到 50Hz", "最后一阶不超过 50Hz 时，需提高 MT 重跑"],
            ["MOTAI-*.PNG/BMP", "模态图", "附录A 模态振型图", "缺图时报告附录不完整"],
            ["SQ-*", "方钢应力云图", "附录B 方钢应力图", "缺失时附录B不完整"],
            ["TB*/TD*", "托臂应力云图", "方钢外边长 <=120mm 时附录C", "该有没有、或不该有却输出，说明截面分支错"],
            ["UNKNOWN", "软件没识别出节点/关键点/集合名", "正式表格中不得出现", "必须回查建模集合、PIP后处理集合和 LIS 原始行"],
            ["source_ref", "每个数从哪里来的说明", "追溯 Excel、LIS、命令流、规范或输入文件", "没有 source_ref 的公式/数据不能作为正式结论"],
            ["source_materials", "原始资料库", "标准命令流、PIP、MAC、SECT、Excel、PDF、DOCX", "只读，不允许改；部署包必须包含必要最小标准命令流包"],
        ],
    )

    _add_note(document, "1. 提资解析规则")
    _add_table(
        document,
        ["规则项", "已确认规则", "代码/审查要求"],
        [
            ["字段识别", "自动识别提资中的主体字段，不按固定列号；I列/埋板列本质是方钢截面，但也必须按表头/内容自动识别。", "intake_excel_reader/job_input_builder 必须支持不同表头命名。"],
            ["报告号/计算批次", "初版新提资可能没有报告号；有报告号显示报告号，没有则显示提资单号或临时批次，后续可补号替换。", "不得要求用户手填行号才能解析。"],
            ["层数", "识别几层就建几层；例如 18185NI-LXSJ4140 是双侧2+2层，不得误判为3层。", "解析结果写入 tray_load_mapping，并在网页显示。"],
            ["三侧", "提资出现三侧时先解析并提示，但生产计算暂阻断，输出“需要人工确认”的命令流，直到三侧标准命令流和后处理映射被审定。", "禁止用单/双侧模板硬算三侧。"],
            ["托盘载荷", "解析托盘宽度、层数、托臂长度、等效密度；无法识别时提示人工复核。", "不能全零或默认乱填。"],
            ["高/低侧梆", "没有动力电缆类型、且提资 Excel 未标明高低侧梆时，默认高侧梆。", "默认值必须写入 audit。"],
        ],
    )

    _add_note(document, "2. 分析方法和反应谱规则")
    _add_table(
        document,
        ["规则项", "已确认规则", "门禁"],
        [
            ["钢平台", "只看是否明确为钢平台；钢平台采用静力法。", "钢平台仍可需要谱文件，用于等效静力加速度。"],
            ["非钢平台", "未显示钢平台的 S2 支架采用反应谱法。", "反应谱文件按项目选择，不能固定某个样例谱。"],
            ["阻尼比", "反应谱阻尼比采用 SL-1=7%，SL-2=10%。", "选谱输出 spectrum_selection.json。"],
            ["谱表匹配", "按项目代号、厂房/谱表、标高、方向匹配；只看当前项目名称代号，不混用其他项目同名厂房。", "缺谱或项目不匹配必须 fail。"],
            ["标高", "不采用历史报告中的多标高包络；使用当前提资标高。", "若报告写了多标高但提资只给单标高，记录为历史报告/源文件冲突。"],
            ["静力修正", "paox/paoy/paoz/pasx/pasy/pasz 来自所选谱 100Hz 加速度，不取负号。", "100Hz 不存在时提示/阻断；不得用负号。"],
            ["载荷展示", "基础载荷和连接螺栓载荷取绝对值后的包络值，报告表中不出现负号。", "保留 raw 值和包络值来源。"],
        ],
    )

    _add_note(document, "3. APDL 命令流和 MT 规则")
    _add_table(
        document,
        ["规则项", "已确认规则", "门禁"],
        [
            ["三份命令流", "每个 job 输出 generated_model.mac、generated_solve.mac、generated_post.mac；run_all.mac 只按顺序调用三份命令。", "便于科室按建模/计算/提取审核。"],
            ["命令来源", "必须基于标准命令流/PIP/MAC/SECT 参数化生成，不从零瞎写。", "source_ref 记录源文件和规则。"],
            ["MT 初值", "求解前取提资几何规则、已审源命令流 MT、默认下限三者中的较大值；当前几何规则为 MT=15*托盘总层数。", "MT 是求解前指定值，不能从 Mode.oup 倒推本次求解。"],
            ["MT 覆盖校核", "运行后读取 Mode.oup，最后一阶频率必须超过 50Hz。", "未超过 50Hz 时门禁 fail，提高 MT 重跑。"],
            ["静力修正", "计算命令流中必须含静力修正过程，100Hz 加速度不取负号。", "门禁扫描 paox/paoy/paoz/pasx/pasy/pasz。"],
        ],
    )

    _add_note(document, "4. 材料与许用应力规则")
    _add_table(
        document,
        ["规则项", "已确认规则", "source_ref/门禁"],
        [
            ["默认材料", "非钢平台默认 Q355；钢平台中只有方钢支架按 Q235 保守评定，其它构件按构件策略使用 Q355。", "材料策略写入 material_policy_metadata。"],
            ["Q355 许用", "Q355 弯曲许用值 234.30 MPa，剪切 142 MPa，拉伸按新版 Excel 表。", "电缆桥架--.xlsx:Q355。"],
            ["Q235 许用", "Q235 弯曲许用值 155.10 MPa，剪切 94 MPa，拉伸按新版 Excel 表。", "电缆桥架--.xlsx:Q235。"],
            ["压缩许用", "A、I 来自“许用应力”表；L 来自当前方钢长度；计算 KL/R 和 Cc 后，KL/R<Cc 用方程4，否则用方程5。", "source_ref 必须含 A/I 单元格、L、KL/R、Cc、方程编号。"],
            ["组合应力", "拉弯组合、压弯组合的许用值均为 1。", "按报告表约定输出 ratio sum。"],
            ["Python/Excel", "Python 复刻结果应与 Excel 权威结果一致；若冲突则 fail，不能强行通过。", "Excel COM 不可用时，未确认公式不得自动 pass。"],
        ],
    )

    _add_note(document, "5. 方钢截面、焊缝和附录C分支")
    _add_table(
        document,
        ["截面条件", "应输出内容", "不得输出/注意"],
        [
            ["方钢外边长 <=120mm（含 120*120*6、120*120*8、120*120*10）", "表6-2-1 托臂根部焊缝评定结果（应力比）；等效应力系数 0.526；附录C 托臂应力图 TB*/TD*。", "不输出大截面用的表6-2 托臂根部所受载荷；不强制依赖 HF-FORCE.LIS。"],
            ["方钢外边长 >120mm", "表6-2 托臂根部所受载荷；表6-3 托臂根部焊缝评定结果；附录C 焊缝评定原理。", "不输出托臂云图；HF-FORCE.LIS 缺失则阻断。"],
            ["提资缺方钢截面", "试算候选 SECT，选择 ratio<1 且最接近 1 的截面；若 100*100*6 已远小于1，也可选 100*100*6。", "不能选 ratio=1 或 >1；候选结果要给用户看。"],
            ["待继续确认", "小截面表6-2-1 是否需要组合项已被标注“需要组合项”，当前代码应保留组合项门禁。", "如果科室确认表格项变化，需要同步修改 chapter6_display、template_injector 和 golden test。"],
        ],
    )

    _add_note(document, "6. 结果提取映射")
    _add_table(
        document,
        ["结果类别", "正式来源", "必须检查"],
        [
            ["支架方钢应力", "MAXBEAMSTRESS.LIS / SQUAREBEAMSTRESS.LIS 中支架方钢集合", "不得与托臂 TMAX 混用；不能按“谁更接近报告”选择。"],
            ["托臂应力", "TMAXBEAMSTRESS.LIS 中托臂集合", "全零时必须查集合、工况和 PIP 后处理命令。"],
            ["基础载荷", "JCZH.LIS", "工况 DW、SL-1、SL-2；FX/FY/FZ/MY/MZ；不得出现 UNKNOWN。"],
            ["连接螺栓载荷", "LS-FORCE.LIS 中托盘与托臂固定点集合", "UPSET/FAULTED 两工况；节点号/关键点必须可追溯。"],
            ["大截面焊缝载荷", "HF-FORCE.LIS", "只用于 >120mm 分支；缺失时不能填空白表。"],
            ["模态", "Mode.oup + MOTAI 图片", "频率表和图均要提取；MT 覆盖 50Hz。"],
        ],
    )

    _add_note(document, "7. 报告注入规则")
    _add_table(
        document,
        ["章节/图表", "规则", "门禁"],
        [
            ["表3-1 反应谱", "按所选项目、厂房/谱表、标高替换表头和数据；不要只截取前几十行。", "保持模板格式不变。"],
            ["图5.1", "S2 支架有限元模型图。", "必须来自 ANSYS 输出或按已审集合导出，不用错误节点图替代。"],
            ["图5.2", "托臂有限元模型图。", "必须与模板/报告结构一致。"],
            ["第六章", "表格标题、顺序、该有/不该有的表必须按截面分支和模板类型输出。", "不得出现空白评定表或全零表作为正式结论。"],
            ["附录A", "模态频率和模态图。", "缺图则报告 audit fail。"],
            ["附录B", "方钢应力图。", "按 SQ 图输出。"],
            ["附录C", "<=120mm 输出托臂应力图；>120mm 输出焊缝评定原理。", "两种模式不能混出。"],
        ],
    )

    _add_note(document, "8. AI 与部署边界")
    _add_table(
        document,
        ["项目", "规则"],
        [
            ["AI 角色", "AI 做提资理解辅助、命令流模板编排建议、日志解释、异常质控和修复建议；正式计算仍由 ANSYS、Excel/确定性公式和 source_ref 决定。"],
            ["大模型失败", "不能因为大模型超时就假装通过；可以退回规则性质控，但页面必须说明审查来源。"],
            ["部署包", "必须包含必要最小标准命令流包，否则会出现 “No APDL/PIP model families found”。部署不能修改 source_materials 原件。"],
            ["内网访问", "非白名单 IP 不进入平台，应提示“请联系管理员-duxyb”。"],
        ],
    )

    _add_note(document, "9. 当前仍需科室确认的问题")
    _add_table(
        document,
        ["编号", "问题", "当前处理"],
        [
            ["Q1", "MT 初值是否还需按托盘长度、支架高度、单/双侧分档？", "当前按 max(15*层数, 源 MT, 默认下限)；运行后用 50Hz 门禁校核。"],
            ["Q2", "小截面表6-2-1 的组合项最终格式。", "按批注“需要组合项”保留组合项。"],
            ["Q3", "三侧 S2 是否允许由参数化模板直接计算。", "当前解析但阻断生产计算，需三侧标准命令流/后处理集合审定。"],
            ["Q4", "历史报告若与标准命令流/提资/谱选择冲突，是否作为报告问题单独留存。", "当前不硬凑，记录冲突，不作为新提资生产逻辑依据。"],
            ["Q5", "图5.1/图5.2 的 ANSYS 输出视角和集合命名是否需统一模板。", "应按模板图号和对应集合导出，缺失时报告 audit fail。"],
        ],
    )

    for section in document.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = "Microsoft YaHei"
                        run.font.size = Pt(9)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


if __name__ == "__main__":
    build_checklist(Path("docs/S2_RULE_CONFIRMATION_CHECKLIST.docx"))
