from __future__ import annotations

import json
import math
import re
from pathlib import Path
from typing import Any

from core.intake.tray_load_parser import parse_tray_load_description
from core.optimizer.square_section_selector import parse_square_section_name


def _number(text: str | None) -> float | None:
    if not text:
        return None
    match = re.search(r"[-+]?(?:\d+(?:\.\d*)?|\.\d+)", str(text))
    return float(match.group(0)) if match else None


def _normalise_section(value: Any) -> str | None:
    if value is None:
        return None
    text = str(value).replace("*", "-").replace("×", "-").replace("x", "-").replace("X", "-")
    match = re.search(r"(\d+(?:\.\d+)?)\s*-\s*(\d+(?:\.\d+)?)\s*-\s*(\d+(?:\.\d+)?)", text)
    if not match:
        return None
    return "-".join(str(int(float(part))) for part in match.groups())


def _report_text(report_path: Path | str) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx is required to read baseline reports") from exc
    document = Document(str(report_path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.split())
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = " ".join(cell.text.split())
                if text:
                    parts.append(text)
    return "\n".join(parts)


def _unique_floats(values: list[float], *, ndigits: int = 6) -> list[float]:
    unique: list[float] = []
    seen: set[float] = set()
    for value in values:
        rounded = round(float(value), ndigits)
        if rounded not in seen:
            unique.append(float(value))
            seen.add(rounded)
    return unique


def _read_source_text(path: Path) -> str:
    data = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gbk", "gb2312"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _extract_envelope_elevations_from_line(line: str) -> list[float]:
    if "Envelop" not in line and "Envelope" not in line and "包络" not in line:
        return []
    values: list[float] = []
    for match in re.finditer(r"\([^)]*?,\s*([-+]?\d+(?:\.\d+)?)\s*\)", line):
        value = float(match.group(1))
        if -100.0 <= value <= 300.0:
            values.append(value)
    if values:
        return _unique_floats(values)
    if "m" in line:
        for match in re.finditer(r"([-+]?\d+(?:\.\d+)?)\s*m", line, flags=re.IGNORECASE):
            value = float(match.group(1))
            if -100.0 <= value <= 300.0:
                values.append(value)
    return _unique_floats(values)


def _extract_static_filename_elevations(path: Path) -> list[float]:
    name = path.stem
    if "静力" not in name and "计算文件" not in name:
        return []
    values: list[float] = []
    for match in re.finditer(r"(?<!\d)([-+]?\d+\.\d+)(?!\d)", name):
        value = float(match.group(1))
        if -100.0 <= value <= 300.0:
            values.append(value)
    return _unique_floats(values)


def _extract_command_spectrum_features(report_path: Path | str) -> dict[str, Any]:
    report_path = Path(report_path)
    calc_dir = report_path.parent / "计算文件"
    if not calc_dir.exists():
        return {
            "command_spectrum_elevations": [],
            "command_spectrum_source_refs": [],
            "command_spectrum_source": None,
        }
    source_files = [
        path
        for path in sorted(calc_dir.rglob("*"))
        if path.is_file() and path.suffix.lower() in {".mac", ".pip", ".txt"}
    ]
    line_elevations: list[float] = []
    source_refs: list[dict[str, Any]] = []
    for source_file in source_files:
        text = _read_source_text(source_file)
        for line_number, line in enumerate(text.splitlines(), start=1):
            elevations = _extract_envelope_elevations_from_line(line)
            if not elevations:
                continue
            line_elevations.extend(elevations)
            source_refs.append(
                {
                    "file": str(source_file),
                    "line": line_number,
                    "text": line.strip(),
                    "elevations": elevations,
                    "evidence": "command_envelope_line",
                }
            )
    if line_elevations:
        return {
            "command_spectrum_elevations": _unique_floats(line_elevations),
            "command_spectrum_source_refs": source_refs,
            "command_spectrum_source": "command_envelope_line",
        }

    filename_elevations: list[float] = []
    filename_refs: list[dict[str, Any]] = []
    for source_file in source_files:
        elevations = _extract_static_filename_elevations(source_file)
        if not elevations:
            continue
        filename_elevations.extend(elevations)
        filename_refs.append(
            {
                "file": str(source_file),
                "line": None,
                "text": source_file.name,
                "elevations": elevations,
                "evidence": "static_calculation_filename",
            }
        )
    return {
        "command_spectrum_elevations": _unique_floats(filename_elevations),
        "command_spectrum_source_refs": filename_refs,
        "command_spectrum_source": "static_calculation_filename" if filename_elevations else None,
    }


def extract_report_design_features(report_path: Path | str) -> dict[str, Any]:
    """Extract input-side design facts from the report narrative.

    This is used only to choose the matching intake row for a report number.
    It never looks at calculated values, so it cannot tune the model to the
    baseline by numerical closeness.
    """

    text = _report_text(report_path)
    static_spectrum_elevations: list[float] = []
    for sentence in re.split(r"(?<=[。；;])\s*", text):
        if "反应谱" not in sentence or "m" not in sentence:
            continue
        if "包络" not in sentence and "楼层" not in sentence:
            continue
        for match in re.finditer(r"([-+]?\d+(?:\.\d+)?)\s*m", sentence, flags=re.IGNORECASE):
            try:
                value = float(match.group(1))
            except ValueError:
                continue
            if -100.0 <= value <= 300.0 and value not in static_spectrum_elevations:
                static_spectrum_elevations.append(value)
    design_sentence = ""
    for sentence in re.split(r"(?<=[。；;])\s*", text):
        if "支架方钢长度" in sentence and "托盘" in sentence:
            design_sentence = sentence
            break
    if not design_sentence:
        for sentence in re.split(r"(?<=[。；;])\s*", text):
            if "截面大小" in sentence and "最大间距" in sentence:
                design_sentence = sentence
                break
    section_match = re.search(r"截面大小为\s*(\d+(?:\.\d+)?)\s*[*×xX-]\s*(\d+(?:\.\d+)?)\s*[*×xX-]\s*(\d+(?:\.\d+)?)\s*mm", design_sentence)
    side_match = re.search(r"(单侧|双侧|两侧|三侧)\s*(\d+(?:\s*\+\s*\d+){0,2})\s*层", design_sentence)
    width_matches = [
        int(float(match.group(1)))
        for match in re.finditer(r"(\d+(?:\.\d+)?)\s*mm\s*托盘", design_sentence)
    ]
    load_matches = [
        float(match.group(1))
        for match in re.finditer(r"载荷为\s*(\d+(?:\.\d+)?)\s*kg/m", design_sentence, flags=re.IGNORECASE)
    ]
    height_match = re.search(r"支架方钢长度为\s*(\d+(?:\.\d+)?)\s*m", design_sentence, flags=re.IGNORECASE)
    spacing_match = re.search(r"最大间距为\s*(\d+(?:\.\d+)?)\s*m", design_sentence, flags=re.IGNORECASE)
    command_spectrum = _extract_command_spectrum_features(report_path)
    command_elevations = command_spectrum.get("command_spectrum_elevations") or []
    report_command_conflict = bool(
        command_elevations
        and static_spectrum_elevations
        and set(round(float(item), 3) for item in command_elevations)
        != set(round(float(item), 3) for item in static_spectrum_elevations)
    )
    selected_spectrum_elevations = command_elevations or static_spectrum_elevations
    features = {
        "report_path": str(report_path),
        "design_sentence": design_sentence,
        "support_height_m": _number(height_match.group(1) if height_match else None),
        "support_spacing_m": _number(spacing_match.group(1) if spacing_match else None),
        "square_section_spec": None,
        "side_word": side_match.group(1) if side_match else None,
        "layer_counts": [int(item) for item in side_match.group(2).replace(" ", "").split("+")] if side_match else [],
        "tray_width_mm": width_matches[0] if width_matches else None,
        "tray_widths_mm": sorted(set(width_matches)),
        "tray_load_kg_per_m": load_matches[0] if load_matches else None,
        "tray_loads_kg_per_m": sorted(set(round(item, 6) for item in load_matches)),
        "static_spectrum_elevations": selected_spectrum_elevations,
        "report_spectrum_elevations": static_spectrum_elevations,
        "command_spectrum_elevations": command_elevations,
        "command_spectrum_source_refs": command_spectrum.get("command_spectrum_source_refs") or [],
        "spectrum_elevation_source": command_spectrum.get("command_spectrum_source") or ("report_text" if static_spectrum_elevations else None),
        "report_command_spectrum_conflict": report_command_conflict,
    }
    if section_match:
        features["square_section_spec"] = "-".join(str(int(float(part))) for part in section_match.groups())
    return features


def _tray_features(description: Any) -> dict[str, Any]:
    try:
        parsed = parse_tray_load_description(description or "")
    except Exception as exc:
        return {"status": "fail", "error": str(exc), "description": str(description or "")}
    layers = parsed.get("layers") or []
    widths = [int(layer.get("tray_width_mm")) for layer in layers if layer.get("tray_width_mm")]
    loads = [float(layer.get("load_kg_per_m")) for layer in layers if layer.get("load_kg_per_m") is not None]
    return {
        "status": "pass",
        "front_layers": int(parsed.get("front_layers") or 0),
        "back_layers": int(parsed.get("back_layers") or 0),
        "declared_layers": parsed.get("declared_layers") or [],
        "widths": widths,
        "unique_widths": sorted(set(widths)),
        "loads": loads,
        "unique_loads": sorted(set(round(load, 6) for load in loads)),
    }


def _float_match(a: Any, b: Any, *, tol: float = 1e-6) -> bool:
    if a is None or b is None:
        return False
    return math.isclose(float(a), float(b), rel_tol=tol, abs_tol=tol)


def _side_count(side_word: str | None) -> int | None:
    if not side_word:
        return None
    if side_word == "单侧":
        return 1
    if side_word in {"双侧", "两侧"}:
        return 2
    if side_word == "三侧":
        return 3
    return None


def score_intake_row_against_report(row: dict[str, Any], report_features: dict[str, Any]) -> dict[str, Any]:
    tray = _tray_features(row.get("description"))
    score = 0
    checks: list[dict[str, Any]] = []

    def add(check_id: str, passed: bool, weight: int, value: Any, expected: Any, source: Any = None) -> None:
        nonlocal score
        if passed:
            score += weight
        check = {
            "check_id": check_id,
            "status": "pass" if passed else "fail",
            "weight": weight,
            "value": value,
            "expected": expected,
        }
        if source is not None:
            check["source"] = source
        checks.append(check)

    add("support_height_m", _float_match(row.get("support_height_m"), report_features.get("support_height_m")), 4, row.get("support_height_m"), report_features.get("support_height_m"))
    add("support_spacing_m", _float_match(row.get("support_spacing_m"), report_features.get("support_spacing_m")), 4, row.get("support_spacing_m"), report_features.get("support_spacing_m"))
    add("square_section_spec", _normalise_section(row.get("square_section_spec")) == report_features.get("square_section_spec"), 4, _normalise_section(row.get("square_section_spec")), report_features.get("square_section_spec"))
    expected_elevations = [float(item) for item in (report_features.get("static_spectrum_elevations") or [])]
    if expected_elevations:
        actual_elevations: list[float] = []
        for item in row.get("elevation_candidates") or []:
            if item is not None:
                actual_elevations.append(float(item))
        if row.get("elevation") is not None:
            actual_elevations.append(float(row.get("elevation")))
        actual_elevations = sorted(set(round(item, 6) for item in actual_elevations))
        add(
            "spectrum_elevation_m",
            all(
                any(_float_match(actual, expected, tol=1e-3) for actual in actual_elevations)
                for expected in expected_elevations
            ),
            4,
            actual_elevations,
            expected_elevations,
            {
                "source": report_features.get("spectrum_elevation_source"),
                "source_refs": report_features.get("command_spectrum_source_refs") or [],
                "report_spectrum_elevations": report_features.get("report_spectrum_elevations") or [],
                "report_command_spectrum_conflict": bool(report_features.get("report_command_spectrum_conflict")),
            },
        )
    if tray.get("status") == "pass":
        expected_width = report_features.get("tray_width_mm")
        add("tray_width_mm", expected_width in tray.get("unique_widths", []), 3, tray.get("unique_widths"), expected_width)
        expected_widths = report_features.get("tray_widths_mm") or ([expected_width] if expected_width is not None else [])
        if expected_widths:
            actual_widths = set(tray.get("unique_widths", []))
            add(
                "tray_widths_mm",
                all(int(width) in actual_widths for width in expected_widths),
                4,
                tray.get("unique_widths"),
                expected_widths,
            )
        expected_layers = report_features.get("layer_counts") or []
        if _side_count(report_features.get("side_word")) == 1 and expected_layers:
            add("single_side_layer_count", int(tray.get("front_layers") or 0) == int(expected_layers[0]), 3, tray.get("front_layers"), expected_layers[0])
        elif _side_count(report_features.get("side_word")) == 2 and expected_layers:
            actual = [int(tray.get("front_layers") or 0), int(tray.get("back_layers") or 0)]
            expected = expected_layers if len(expected_layers) > 1 else expected_layers * 2
            add("double_side_layer_count", actual == expected, 3, actual, expected)
        elif _side_count(report_features.get("side_word")) == 3 and expected_layers:
            actual = [int(item) for item in (tray.get("declared_layers") or [])]
            add("three_side_layer_count", actual == expected_layers, 4, actual, expected_layers)
        expected_load = report_features.get("tray_load_kg_per_m")
        if expected_load is not None:
            add(
                "tray_load_kg_per_m",
                any(_float_match(load, expected_load, tol=1e-3) for load in tray.get("unique_loads", [])),
                2,
                tray.get("unique_loads"),
                expected_load,
            )
        expected_loads = report_features.get("tray_loads_kg_per_m") or ([expected_load] if expected_load is not None else [])
        if expected_loads:
            actual_loads = tray.get("unique_loads", [])
            add(
                "tray_loads_kg_per_m",
                all(any(_float_match(actual, expected, tol=1e-3) for actual in actual_loads) for expected in expected_loads),
                3,
                actual_loads,
                expected_loads,
            )
    else:
        add("tray_parse", False, 5, tray.get("error"), "parseable tray description")

    required_check_ids = {
        "support_height_m",
        "support_spacing_m",
        "square_section_spec",
        "tray_width_mm",
        "tray_widths_mm",
        "single_side_layer_count",
        "double_side_layer_count",
        "three_side_layer_count",
        "tray_load_kg_per_m",
        "tray_loads_kg_per_m",
        "spectrum_elevation_m",
    }
    failed_required_checks = [
        check
        for check in checks
        if check["check_id"] in required_check_ids and check["status"] != "pass"
    ]
    return {
        "score": score,
        "row_number": row.get("intake_row_number"),
        "intake_serial": row.get("intake_serial"),
        "report_number": row.get("report_number"),
        "description": row.get("description"),
        "checks": checks,
        "required_check_count": len([check for check in checks if check["check_id"] in required_check_ids]),
        "failed_required_check_count": len(failed_required_checks),
        "failed_required_checks": failed_required_checks,
        "tray_features": tray,
    }


def select_representative_intake_rows(rows: list[dict[str, Any]], reports_by_number: dict[str, str]) -> dict[str, Any]:
    grouped: dict[str, list[dict[str, Any]]] = {}
    for row in rows:
        report_no = str(row.get("report_number") or row.get("calculation_batch") or "").strip()
        if report_no in reports_by_number:
            grouped.setdefault(report_no, []).append(row)

    selected: list[dict[str, Any]] = []
    blocked: list[dict[str, Any]] = []
    for report_no in sorted(reports_by_number):
        candidates = grouped.get(report_no) or []
        if not candidates:
            blocked.append({"report_no": report_no, "status": "blocked", "reason": "no intake row for report"})
            continue
        report_features = extract_report_design_features(reports_by_number[report_no])
        scored = [score_intake_row_against_report(row, report_features) for row in candidates]
        scored.sort(key=lambda item: (-int(item["score"]), int(item.get("row_number") or 10**9)))
        best = scored[0]
        if best["score"] <= 0:
            blocked.append(
                {
                    "report_no": report_no,
                    "status": "blocked",
                    "reason": "report design features could not be matched to intake rows",
                    "report_features": report_features,
                    "candidate_scores": scored,
                }
            )
            continue
        matches = [item for item in scored if item["score"] == best["score"]]
        report_command_spectrum_conflict = bool(report_features.get("report_command_spectrum_conflict"))
        selection_has_conflict = bool(best.get("failed_required_check_count")) or report_command_spectrum_conflict
        if best.get("failed_required_check_count"):
            conflict_reason = (
                "No intake row fully matches the report design facts; this report-number validation case is "
                "a historical intake/report conflict and must not be used to tune production logic."
            )
        elif report_command_spectrum_conflict:
            conflict_reason = (
                "The report narrative spectrum elevations conflict with the same report package's command-file "
                "spectrum elevations; this historical report/source conflict must be reviewed instead of tuning "
                "production spectrum selection."
            )
        else:
            conflict_reason = None
        selected.append(
            {
                "report_no": report_no,
                "intake_row_number": best.get("row_number"),
                "intake_serial": best.get("intake_serial"),
                "baseline_report": reports_by_number[report_no],
                "score": best["score"],
                "tie_count": len(matches),
                "report_features": report_features,
                "selected_score": best,
                "candidate_scores": scored,
                "selection_conflict_status": "conflict" if selection_has_conflict else "pass",
                "selection_conflict_reason": conflict_reason,
                "failed_required_checks": best.get("failed_required_checks") or [],
                "selection_policy": "Select by report design facts versus intake fields only; ANSYS results and numerical closeness are not used.",
            }
        )

    return {
        "status": "pass" if selected and not blocked else "blocked",
        "selected_count": len(selected),
        "blocked_count": len(blocked),
        "selected": selected,
        "blocked": blocked,
        "policy": "Report-number validation uses one representative intake row per baseline report, selected from report narrative design features and intake fields.",
    }


def write_selection_audit(path: Path | str, payload: dict[str, Any]) -> None:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
