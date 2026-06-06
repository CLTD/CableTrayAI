from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from core.validation.analysis_scope import (
    SQUARE_SECTION_CANTILEVER_CLOUD_MAX_OUTER_MM,
    SQUARE_SECTION_WELD_PRINCIPLE_GT_OUTER_MM,
    appendix_c_mode_for_square_section,
    parse_square_outer_width_mm,
    parse_square_section_spec,
)


def _docx_text(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def square_section_from_report_text(text: str) -> dict[str, Any]:
    patterns = [
        r"截面大小为\s*([0-9.]+\s*[*xX×-]\s*[0-9.]+\s*[*xX×-]\s*[0-9.]+)\s*mm",
        r"方钢.*?([0-9.]+\s*[*xX×-]\s*[0-9.]+\s*[*xX×-]\s*[0-9.]+)\s*mm",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            raw = match.group(1)
            parsed = parse_square_section_spec(raw) or {"outer_mm": parse_square_outer_width_mm(raw), "thickness_mm": None}
            return {"raw": raw, "outer_mm": parsed.get("outer_mm"), "thickness_mm": parsed.get("thickness_mm")}
    return {"raw": None, "outer_mm": None, "thickness_mm": None}


def expected_appendix_c_mode(square_outer_width_mm: float | None) -> str:
    return appendix_c_mode_for_square_section(square_outer_width_mm)[0]


def observed_appendix_c_mode(text: str) -> str:
    if "附录C：托臂应力图" in text or "附录C:托臂应力图" in text:
        return "cantilever_stress_cloud"
    if "附录C：焊缝评定原理" in text or "附录C:焊缝评定原理" in text:
        return "weld_evaluation_principle"
    if "附录C" in text:
        return "appendix_c_other"
    return "none"


def audit_report_section_appendix_rule(report_path: Path | str) -> dict[str, Any]:
    report_path = Path(report_path)
    text = _docx_text(report_path)
    section = square_section_from_report_text(text)
    expected = expected_appendix_c_mode(section.get("outer_mm"))
    observed = observed_appendix_c_mode(text)
    if expected == "needs_square_section_selection":
        status = "needs_review"
    elif expected == observed:
        status = "pass"
    else:
        status = "source_conflict"
    return {
        "status": status,
        "report": report_path.stem,
        "report_path": str(report_path),
        "square_section_raw": section.get("raw"),
        "square_outer_width_mm": section.get("outer_mm"),
        "square_thickness_mm": section.get("thickness_mm"),
        "cantilever_cloud_max_outer_mm": SQUARE_SECTION_CANTILEVER_CLOUD_MAX_OUTER_MM,
        "weld_principle_gt_outer_mm": SQUARE_SECTION_WELD_PRINCIPLE_GT_OUTER_MM,
        "expected_appendix_c_mode": expected,
        "observed_appendix_c_mode": observed,
        "rule": "square outer width <= 120 -> cantilever stress cloud; > 120 -> weld evaluation principle",
    }


def audit_report_tree(
    source_root: Path | str = Path("source_materials/model_commands"),
    *,
    output_json: Path | str = Path("docs/SECTION_APPENDIX_RULE_AUDIT.json"),
    output_md: Path | str = Path("docs/SECTION_APPENDIX_RULE_AUDIT.md"),
) -> dict[str, Any]:
    source_root = Path(source_root)
    rows = []
    for report in sorted(source_root.rglob("*.docx"), key=lambda item: item.as_posix().lower()):
        if report.name.startswith("~$"):
            continue
        rows.append(audit_report_section_appendix_rule(report))
    payload = {
        "status": "pass" if rows and all(row["status"] == "pass" for row in rows) else "needs_review",
        "source_root": str(source_root),
        "report_count": len(rows),
        "conflict_count": len([row for row in rows if row["status"] == "source_conflict"]),
        "needs_review_count": len([row for row in rows if row["status"] == "needs_review"]),
        "rows": rows,
    }
    output_json = Path(output_json)
    output_json.parent.mkdir(parents=True, exist_ok=True)
    output_json.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = [
        "# Section Appendix Rule Audit",
        "",
        f"Status: `{payload['status']}`",
        f"Reports: {payload['report_count']}",
        f"Conflicts: {payload['conflict_count']}",
        "",
        "| Report | Square outer | Thickness | Expected appendix C | Observed appendix C | Status |",
        "| --- | ---: | ---: | --- | --- | --- |",
    ]
    for row in rows:
        lines.append(
            f"| {row['report']} | {row.get('square_outer_width_mm') or ''} | {row.get('square_thickness_mm') or ''} | {row['expected_appendix_c_mode']} | {row['observed_appendix_c_mode']} | {row['status']} |"
        )
    Path(output_md).write_text("\n".join(lines) + "\n", encoding="utf-8")
    return payload
