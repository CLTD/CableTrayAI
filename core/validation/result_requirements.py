from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any


MODEL_FIGURES = ["SHITI.PNG", "TBMODEL.PNG"]
MODAL_FIGURES = [f"MOTAI-{index}.PNG" for index in range(1, 5)]
MAIN_STRESS_FIGURES = [
    "B1SDIR1.PNG",
    "B2SDIR2.PNG",
    "B3SBEND.PNG",
    "B4SHEAR.PNG",
    "D1SDIR1.PNG",
    "D2SDIR2.PNG",
    "D3SBEND.PNG",
    "D4SHEAR.PNG",
]
CANTILEVER_FIGURES = [
    "TB1SDIR1.PNG",
    "TB2SDIR2.PNG",
    "TB3SBEND.PNG",
    "TB4SHEAR.PNG",
    "TD1SDIR1.PNG",
    "TD2SDIR2.PNG",
    "TD3SBEND.PNG",
    "TD4SHEAR.PNG",
]


def _safe_text_from_docx(report_path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx is required to classify report requirements") from exc
    document = Document(str(report_path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def classify_report_requirements(report_path: Path | str) -> dict[str, Any]:
    """Classify what should be calculated/extracted for a report.

    This is intentionally caption/section driven.  It avoids the earlier
    mistake of exporting the same appendix C/cantilever figures for cases whose
    report does not contain appendix C.
    """

    report_path = Path(report_path)
    text = _safe_text_from_docx(report_path)
    has_appendix_a = "附录A" in text and "模态" in text
    has_appendix_b = "图B-1" in text or "附录B" in text
    has_appendix_c = "附录C：托臂应力图" in text or "附录C:托臂应力图" in text
    has_cantilever_eval = bool(re.search(r"托臂.*应力评定|支架托臂应力评定", text))
    has_cantilever_weld = "托臂根部焊缝" in text
    has_tray_arm_bolt = "托盘和托臂" in text or "托盘与托臂" in text
    mentions_steel_platform = "钢平台" in text
    explicit_static = "静力法" in text or "采用静力" in text
    uses_spectrum = "谱分析" in text or "反应谱" in text or "楼层反应谱" in text
    analysis_method = "static" if mentions_steel_platform or explicit_static else "response_spectrum"
    method_conflict = bool(mentions_steel_platform and uses_spectrum)
    required_figures = list(MODEL_FIGURES)
    if has_appendix_a:
        required_figures.extend(MODAL_FIGURES)
    if has_appendix_b:
        required_figures.extend(MAIN_STRESS_FIGURES)
    if has_appendix_c:
        required_figures.extend(CANTILEVER_FIGURES)

    payload = {
        "status": "pass",
        "report_path": str(report_path),
        "report_number": report_path.stem,
        "classification": "steel_platform" if mentions_steel_platform else "non_steel_platform",
        "analysis_method": analysis_method,
        "mentions_steel_platform": mentions_steel_platform,
        "explicit_static_method": explicit_static,
        "uses_response_spectrum": uses_spectrum,
        "method_conflict_requires_review": method_conflict,
        "has_appendix_a_modal": has_appendix_a,
        "has_appendix_b_main_stress": has_appendix_b,
        "has_appendix_c_cantilever": has_appendix_c,
        "requires": {
            "modal_analysis": bool(has_appendix_a and analysis_method != "static"),
            "modal_figures": has_appendix_a,
            "modal_frequency_table": has_appendix_a,
            "square_support_stress_eval": True,
            "cantilever_stress_eval": has_cantilever_eval,
            "cantilever_root_weld_eval": has_cantilever_weld,
            "foundation_loads": True,
            "tray_arm_connection_loads": has_tray_arm_bolt,
            "bolt_stress_eval": has_tray_arm_bolt,
            "appendix_c_cantilever_figures": has_appendix_c,
        },
        "required_figures": required_figures,
        "forbidden_figures": [] if has_appendix_c else CANTILEVER_FIGURES,
        "policy": [
            "Do not export or publish appendix C/cantilever figures unless the report/intake classification requires appendix C.",
            "Do not use QA stamps, signatures, or non-appendix images as appendix verification figures.",
            "Steel-platform rows and non-steel-platform rows must not share the same calculation/extraction checklist.",
        ],
    }
    return payload


def classify_job_requirements(job_dir: Path | str, report_path: Path | str | None = None) -> dict[str, Any]:
    job_dir = Path(job_dir)
    if report_path:
        payload = classify_report_requirements(report_path)
    else:
        from core.validation.analysis_scope import classify_scope_from_job

        payload = classify_scope_from_job(job_dir)
        payload["report_path"] = None
        payload["report_number"] = payload.get("report_number") or job_dir.name
    (job_dir / "result_requirements.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return payload
