from __future__ import annotations

import json
import math
import re
from pathlib import Path
from typing import Any

from core.validation.analysis_scope import SQUARE_SECTION_CANTILEVER_CLOUD_MAX_OUTER_MM
from core.validation.section_appendix_rule_audit import square_section_from_report_text
from core.validation.manual_baseline import compare_scalar, relative_error


STRICT_TOLERANCE = 0.01
WELD_DEFAULT_COEFFICIENT = 0.526
BASELINE_CONFLICT_FILE = Path("data/calibration/report_baseline_conflicts.json")
PASS_LIKE_STATUSES = {"pass", "baseline_conflict"}


def _cell_text(cell: Any) -> str:
    return " ".join(str(cell.text or "").replace("\xa0", " ").split())


def _number(value: str | float | int | None) -> float | None:
    if value is None:
        return None
    text = str(value).replace(",", "").strip()
    match = re.search(r"[-+]?(?:\d+(?:\.\d*)?|\.\d+)(?:[Ee][-+]?\d+)?", text)
    if not match:
        return None
    try:
        return float(match.group(0))
    except ValueError:
        return None


def _decimal_places(value: str | float | int | None) -> int | None:
    if value is None:
        return None
    text = str(value).replace(",", "").strip()
    match = re.search(r"[-+]?(?:\d+(?:\.\d*)?|\.\d+)(?:[Ee][-+]?\d+)?", text)
    if not match:
        return None
    token = match.group(0)
    if "e" in token.lower():
        return None
    if "." not in token:
        return 0
    return len(token.split(".", 1)[1])


def _contains_all(text: str, words: tuple[str, ...]) -> bool:
    return all(word in text for word in words)


def _rows(table: Any) -> list[list[str]]:
    return [[_cell_text(cell) for cell in row.cells] for row in table.rows]


def _body_tables_with_context(document: Any):
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    elements: list[tuple[str, Any]] = []
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            text = " ".join(Paragraph(child, document).text.split())
            if text:
                elements.append(("paragraph", text))
            continue
        if isinstance(child, CT_Tbl):
            elements.append(("table", Table(child, document)))

    previous_text: list[str] = []
    table_index = 0
    for index, (kind, value) in enumerate(elements):
        if kind == "paragraph":
            previous_text.append(str(value))
            previous_text = previous_text[-8:]
            continue
        table_index += 1
        following_text: list[str] = []
        for next_kind, next_value in elements[index + 1 :]:
            if next_kind == "table":
                break
            following_text.append(str(next_value))
            if len(following_text) >= 8:
                break
        yield table_index, value, list(previous_text), following_text


def _nearest_caption(context: list[str]) -> str:
    for text in reversed(context):
        if text.startswith("表") or re.search(r"表\s*[A-Za-z0-9一二三四五六七八九十\-\.]+", text):
            return text
    return ""


def _weld_coefficient(context: list[str]) -> tuple[float, str]:
    for text in reversed(context):
        if "焊接系数" not in text and "系数" not in text:
            continue
        match = re.search(r"系数为\s*([0-9]+(?:\.[0-9]+)?)", text)
        if match:
            return float(match.group(1)), text
    return WELD_DEFAULT_COEFFICIENT, "default from report wording: weld coefficient must be confirmed"


def _evaluation_mapping(caption: str, flat: str, context: list[str], square_outer_width_mm: float | None = None) -> dict[str, Any]:
    if "托臂根部焊缝" in caption or "焊缝" in caption:
        coefficient, coefficient_ref = _weld_coefficient(context)
        return {
            "evaluation_kind": "weld_equivalent",
            "component": "cantilever_root_weld",
            "result_source_file": "TMAXBEAMSTRESS.LIS",
            "coefficient": coefficient,
            "coefficient_source_ref": coefficient_ref,
            "mapping_source_ref": caption,
        }
    if "托臂" in caption and "应力评定" in caption:
        return {
            "evaluation_kind": "beam_stress",
            "component": "cantilever_arm",
            "result_source_file": "TMAXBEAMSTRESS.LIS",
            "mapping_source_ref": caption,
        }
    if "方钢" in caption and "应力评定" in caption:
        return {
            "evaluation_kind": "beam_stress",
            "component": "square_support",
            "result_source_file": "SQUAREBEAMSTRESS.LIS",
            "mapping_source_ref": caption,
            "scope_note": "Square-support stress tables default to the section-1-only square-support export. A later report-structure pass may promote large-square multi-topology reports to the all-section support export when the report declares a secondary tray-arm bolt-load table.",
        }
    if (
        ("支架应力评定" in caption or ("鏀灦" in caption and "搴斿姏璇勫畾" in caption))
        and "方钢" not in caption
        and "鏂归挗" not in caption
        and "托臂" not in caption
        and "鎵樿噦" not in caption
        and "焊缝" not in caption
        and "鐒婄紳" not in caption
    ):
        return {
            "evaluation_kind": "beam_stress",
            "component": "mixed_beam_type_1",
            "result_source_file": "MAXBEAMSTRESS.LIS",
            "mapping_source_ref": caption,
            "scope_note": "Generic support stress evaluation maps to the all-section support export by report section semantics, not by nearest numeric value.",
        }
    return {
        "evaluation_kind": "unmapped",
        "component": "",
        "result_source_file": None,
        "mapping_source_ref": caption or flat[:80],
    }


def extract_report_baseline(report_path: Path | str) -> dict[str, Any]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx is required to read baseline reports") from exc

    report_path = Path(report_path)
    document = Document(str(report_path))
    full_text_parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            full_text_parts.extend(cell.text for cell in row.cells)
    section_info = square_section_from_report_text("\n".join(full_text_parts))
    square_outer_width_mm = section_info.get("outer_mm")
    baseline: dict[str, Any] = {
        "source_file": str(report_path),
        "square_section": section_info,
        "modal_results": [],
        "evaluation_ratios": [],
        "loads": [],
        "tables_scanned": len(document.tables),
        "mapping_policy": "Report rows are mapped by table caption and source command output file. Numeric nearest-neighbor matching is forbidden.",
    }
    for table_index, table, context, following_context in _body_tables_with_context(document):
        rows = _rows(table)
        flat = " ".join(" ".join(row) for row in rows)
        caption = _nearest_caption(context)
        if ("频" in flat and ("阶数" in flat or "阶 次" in flat or "阶" in flat)):
            baseline["modal_results"].extend(_extract_modal_rows(rows, table_index))
            continue
        if _contains_all(flat, ("工况", "应力比")) and ("计算值" in flat or "等效应力" in flat):
            mapping = _evaluation_mapping(caption, flat, context, square_outer_width_mm)
            baseline["evaluation_ratios"].extend(_extract_evaluation_rows(rows, table_index, caption, mapping))
            continue
        if _contains_all(flat, ("FX", "FY", "FZ")):
            baseline["loads"].extend(_extract_load_rows(rows, table_index, caption, context, following_context))
    _resolve_connection_load_sources(baseline)
    _resolve_square_support_sources(baseline)
    _apply_report_baseline_conflicts(baseline, report_path)
    baseline["status"] = "pass" if baseline["modal_results"] or baseline["evaluation_ratios"] or baseline["loads"] else "blocked"
    return baseline


def _apply_report_baseline_conflicts(baseline: dict[str, Any], report_path: Path) -> None:
    """Attach audited report-baseline conflicts without changing computed values.

    These entries are calibration-only evidence that a historical report row is
    internally inconsistent with another report/source file. They are not used
    for new-intake calculations and they never replace ANSYS results.
    """

    if not BASELINE_CONFLICT_FILE.exists():
        return
    try:
        payload = json.loads(BASELINE_CONFLICT_FILE.read_text(encoding="utf-8-sig"))
    except Exception:
        return
    package_id = report_path.stem
    entries = [
        entry
        for entry in payload.get("conflicts", [])
        if entry.get("enabled", True) and str(entry.get("package_id") or "") == package_id
    ]
    if not entries:
        return
    for collection_name in ("loads", "evaluation_ratios", "modal_results"):
        for row in baseline.get(collection_name, []):
            for entry in entries:
                if _baseline_conflict_entry_matches(row, entry):
                    metrics = entry.get("metrics") or ["*"]
                    row.setdefault("baseline_conflicts", []).append(
                        {
                            "metrics": metrics,
                            "status": "baseline_conflict",
                            "source_ref": entry.get("source_ref"),
                            "reason": entry.get("reason"),
                            "evidence": entry.get("evidence"),
                            "policy": entry.get("policy")
                            or "Exclude only this audited report-baseline conflict from the precision gate; do not alter ANSYS output.",
                        }
                    )


def _baseline_conflict_entry_matches(row: dict[str, Any], entry: dict[str, Any]) -> bool:
    for key in ("load_kind", "evaluation_kind", "component", "stress_type", "table_caption", "source_ref", "case"):
        expected = entry.get(key)
        if expected is not None and str(row.get(key) or "") != str(expected):
            return False
    expected_values = entry.get("values")
    if isinstance(expected_values, dict):
        row_values = row.get("values") or row
        for key, expected in expected_values.items():
            actual = row_values.get(key)
            if actual is None:
                return False
            if abs(float(actual) - float(expected)) > 1e-6:
                return False
    return True


def _baseline_conflict_for_metric(row: dict[str, Any], metric: str) -> dict[str, Any] | None:
    for conflict in row.get("baseline_conflicts") or []:
        metrics = conflict.get("metrics") or ["*"]
        if "*" in metrics or metric in metrics:
            return conflict
    return None


def _resolve_square_support_sources(baseline: dict[str, Any]) -> None:
    """Resolve square-support report tables from report-declared topology, not numeric closeness."""

    square_outer_width_mm = (baseline.get("square_section") or {}).get("outer_mm")
    has_secondary_connection_table = any(
        row.get("load_kind") == "tray_arm_connection_derived_bolt_load"
        or (
            row.get("load_kind") == "tray_arm_connection"
            and isinstance(row.get("bolt_table_width_mm"), (int, float))
            and float(row["bolt_table_width_mm"]) <= 200.0
        )
        for row in baseline.get("loads", [])
    )
    use_all_section_export = (
        isinstance(square_outer_width_mm, (int, float))
        and float(square_outer_width_mm) > SQUARE_SECTION_CANTILEVER_CLOUD_MAX_OUTER_MM
        and has_secondary_connection_table
    )
    for row in baseline.get("evaluation_ratios", []):
        if row.get("component") != "square_support":
            continue
        if use_all_section_export:
            row["result_source_file"] = "MAXBEAMSTRESS.LIS"
            row["scope_note"] = (
                "Large-square support table promoted to the all-section support export because the report declares a "
                "secondary tray-arm bolt-load table. This follows report topology/chapter structure and never selects "
                "the closest numeric result."
            )
        else:
            row["result_source_file"] = "SQUAREBEAMSTRESS.LIS"
            row["scope_note"] = (
                "Square-support table mapped to the section-1-only square-support export because the report does not "
                "declare a secondary tray-arm bolt-load table. This follows report topology/chapter structure and "
                "never selects the closest numeric result."
            )


def _resolve_connection_load_sources(baseline: dict[str, Any]) -> None:
    """Resolve tray-arm load outputs from report topology, not width alone."""

    connection_rows = [
        row
        for row in baseline.get("loads", [])
        if row.get("load_kind") in {"tray_arm_connection", "tray_arm_connection_derived_bolt_load"}
    ]
    has_wide_primary = any(
        row.get("load_kind") == "tray_arm_connection"
        and isinstance(row.get("bolt_table_width_mm"), (int, float))
        and float(row["bolt_table_width_mm"]) > 200.0
        for row in connection_rows
    )
    for row in connection_rows:
        if row.get("load_kind") != "tray_arm_connection_derived_bolt_load":
            continue
        if has_wide_primary:
            row["mapping_policy"] = (
                "Report declares both a wider primary tray-arm connection table and a <=200 mm secondary table; "
                "the <=200 mm table uses topology-preserved LS-FORCE-NODES.LIS."
            )
            continue
        row["load_kind"] = "tray_arm_connection"
        row["result_source_file"] = "LS-FORCE.LIS"
        row["mapping_policy"] = (
            "Report declares only a <=200 mm primary tray-arm connection table; it maps to the standard "
            "LS-FORCE.LIS postprocessor output rather than the secondary topology export."
        )


def _extract_modal_rows(rows: list[list[str]], table_index: int) -> list[dict[str, Any]]:
    extracted: list[dict[str, Any]] = []
    for row_index, row in enumerate(rows, start=1):
        numbers = [_number(cell) for cell in row]
        values = [value for value in numbers if value is not None]
        if len(values) < 2:
            continue
        for pair_index in range(0, len(values) - 1, 2):
            mode = int(values[pair_index])
            frequency = float(values[pair_index + 1])
            if mode <= 0 or frequency <= 0:
                continue
            extracted.append(
                {
                    "mode": mode,
                    "frequency_hz": frequency,
                    "source_ref": f"table {table_index} row {row_index}",
                }
            )
    return extracted


def _header_indices(header: list[str]) -> dict[str, int]:
    indices: dict[str, int] = {}
    for index, name in enumerate(header):
        if "工况" in name:
            indices["case"] = index
        elif "应力类型" in name:
            indices["stress_type"] = index
        elif "等效应力" in name:
            indices["equivalent"] = index
        elif "计算值" in name:
            indices["calculation"] = index
        elif "许用值" in name:
            indices["allowable"] = index
        elif "应力比" in name:
            indices["ratio"] = index
    return indices


def _extract_evaluation_rows(
    rows: list[list[str]],
    table_index: int,
    caption: str,
    mapping: dict[str, Any],
) -> list[dict[str, Any]]:
    extracted: list[dict[str, Any]] = []
    header_index = next((index for index, row in enumerate(rows) if "应力类型" in " ".join(row) and "应力比" in " ".join(row)), None)
    if header_index is None:
        return extracted
    indices = _header_indices(rows[header_index])
    current_case = ""
    for row_index, row in enumerate(rows[header_index + 1 :], start=header_index + 2):
        row_text = " ".join(row)
        if indices.get("case") is not None and row[indices["case"]]:
            current_case = row[indices["case"]]
        stress_type = row[indices["stress_type"]] if indices.get("stress_type") is not None and len(row) > indices["stress_type"] else ""
        if not stress_type:
            stress_type = _first_text(row) or row_text
        calculation_value = _number(row[indices["calculation"]]) if indices.get("calculation") is not None and len(row) > indices["calculation"] else None
        equivalent_value = _number(row[indices["equivalent"]]) if indices.get("equivalent") is not None and len(row) > indices["equivalent"] else None
        allowable_value = _number(row[indices["allowable"]]) if indices.get("allowable") is not None and len(row) > indices["allowable"] else None
        ratio = _number(row[indices["ratio"]]) if indices.get("ratio") is not None and len(row) > indices["ratio"] else None
        if calculation_value is None or allowable_value is None or ratio is None or ratio < 0:
            continue
        extracted.append(
            {
                "case": current_case,
                "stress_type": stress_type,
                "calculation_value": calculation_value,
                "calculation_text": row[indices["calculation"]] if indices.get("calculation") is not None and len(row) > indices["calculation"] else None,
                "calculation_decimals": _decimal_places(row[indices["calculation"]]) if indices.get("calculation") is not None and len(row) > indices["calculation"] else None,
                "equivalent_stress_value": equivalent_value,
                "equivalent_stress_text": row[indices["equivalent"]] if indices.get("equivalent") is not None and len(row) > indices["equivalent"] else None,
                "equivalent_stress_decimals": _decimal_places(row[indices["equivalent"]]) if indices.get("equivalent") is not None and len(row) > indices["equivalent"] else None,
                "allowable_value": allowable_value,
                "ratio": ratio,
                "ratio_text": row[indices["ratio"]] if indices.get("ratio") is not None and len(row) > indices["ratio"] else None,
                "ratio_decimals": _decimal_places(row[indices["ratio"]]) if indices.get("ratio") is not None and len(row) > indices["ratio"] else None,
                "table_caption": caption,
                "source_ref": f"table {table_index} row {row_index}",
                **mapping,
            }
        )
    return extracted


def _load_metric_from_header(cell: str) -> str | None:
    text = cell.upper().replace(" ", "")
    for key in ("FX", "FY", "FZ", "MX", "MY", "MZ"):
        if key in text:
            return key.lower()
    return None


def _bolt_table_width_from_context(context: list[str]) -> tuple[float | None, str | None]:
    for text in context:
        if "mm" not in text.lower():
            continue
        match = re.search(r"(\d+(?:\.\d+)?)\s*mm", text, flags=re.IGNORECASE)
        if match:
            return float(match.group(1)), text
    return None, None


def _load_kind_from_caption(
    caption: str,
    context: list[str],
    following_context: list[str],
    fixed_point: str,
    case: str,
) -> tuple[str, str | None, float | None, str | None]:
    caption_text = caption or ""
    if "托臂根部所受载荷" in caption_text:
        return "cantilever_root_load", "HF-FORCE.LIS", None, None
    if "基础载荷" in caption_text:
        return "foundation", "JCZH.LIS", None, None
    if "支架连接螺栓载荷" in caption_text:
        bolt_width_mm, width_source_ref = _bolt_table_width_from_context(following_context)
        if bolt_width_mm is not None and bolt_width_mm <= 200:
            return "tray_arm_connection_derived_bolt_load", "LS-FORCE-NODES.LIS", bolt_width_mm, width_source_ref
        return "tray_arm_connection", "LS-FORCE.LIS", bolt_width_mm, width_source_ref
    if "托盘" in fixed_point or "托臂" in fixed_point or "托盘" in case or "托臂" in case:
        return "tray_arm_connection", "LS-FORCE.LIS", None, None
    if case.upper() in {"DW", "SL-1", "SL-2"}:
        return "foundation", "JCZH.LIS", None, None
    return "load_table_unmapped", None, None, None


def _extract_load_rows(
    rows: list[list[str]],
    table_index: int,
    caption: str,
    context: list[str],
    following_context: list[str],
) -> list[dict[str, Any]]:
    extracted: list[dict[str, Any]] = []
    header_index = next((index for index, row in enumerate(rows) if any(_load_metric_from_header(cell) for cell in row)), None)
    if header_index is None:
        return extracted
    header = rows[header_index]
    metric_indices = {
        _load_metric_from_header(cell): index
        for index, cell in enumerate(header)
        if _load_metric_from_header(cell)
    }
    case_index = next((index for index, cell in enumerate(header) if "工况" in cell), None)
    fixed_index = next((index for index, cell in enumerate(header) if "固定" in cell), None)
    for row_index, row in enumerate(rows[header_index + 1 :], start=header_index + 2):
        if not row:
            continue
        fixed_point = row[fixed_index] if fixed_index is not None and len(row) > fixed_index else ""
        case = row[case_index] if case_index is not None and len(row) > case_index else (row[0] if row else "")
        values: dict[str, float] = {}
        for metric, index in metric_indices.items():
            if metric is None or len(row) <= index:
                continue
            value = _number(row[index])
            if value is not None:
                values[metric] = value
        if len(values) < 3:
            continue
        load_kind, result_source_file, bolt_width_mm, width_source_ref = _load_kind_from_caption(
            caption,
            context,
            following_context,
            fixed_point,
            case,
        )
        row_payload = {
                "case": case,
                "fixed_point": fixed_point,
                "values": values,
                "units": {"force": "N", "moment": "N*m"},
                "load_kind": load_kind,
                "result_source_file": result_source_file,
                "table_caption": caption,
                "mapping_source_ref": caption,
                "source_ref": f"table {table_index} row {row_index}",
            }
        if bolt_width_mm is not None:
            row_payload["bolt_table_width_mm"] = bolt_width_mm
            row_payload["bolt_table_width_source_ref"] = width_source_ref
            row_payload["mapping_policy"] = (
                "Tray-arm bolt-load tables are classified from the report text immediately following the table: "
                "width <= 200 mm uses the secondary topology export, while larger widths use standard LS-FORCE.LIS."
            )
        extracted.append(row_payload)
    return extracted


def _first_text(values: list[str]) -> str:
    for value in values:
        if value and _number(value) is None:
            return value
    return ""


def compare_report_baseline(result: dict[str, Any], baseline: dict[str, Any], *, tolerance: float = STRICT_TOLERANCE) -> dict[str, Any]:
    comparisons: list[dict[str, Any]] = []
    comparisons.extend(_compare_modal(result, baseline, tolerance))
    comparisons.extend(_compare_beam_calculations(result, baseline, tolerance))
    comparisons.extend(_compare_ratios(result, baseline, tolerance))
    comparisons.extend(_compare_loads(result, baseline, tolerance))
    comparisons.extend(_mapping_blockers(result, baseline))
    if not comparisons:
        return {
            "status": "blocked",
            "precision_verified": False,
            "tolerance": tolerance,
            "message": "No comparable fields were found between result.json and the baseline report.",
            "comparisons": [],
        }
    status = "pass" if all(item["status"] in PASS_LIKE_STATUSES for item in comparisons) else "fail"
    baseline_conflicts = [item for item in comparisons if item.get("status") == "baseline_conflict"]
    finite_relative_errors = [
        item.get("relative_error")
        for item in comparisons
        if isinstance(item.get("relative_error"), (int, float)) and math.isfinite(float(item.get("relative_error")))
    ]
    absolute_errors = [
        item.get("absolute_error")
        for item in comparisons
        if isinstance(item.get("absolute_error"), (int, float))
    ]
    gate_errors = [
        item.get("gate_error", item.get("relative_error"))
        for item in comparisons
        if isinstance(item.get("gate_error", item.get("relative_error")), (int, float))
        and math.isfinite(float(item.get("gate_error", item.get("relative_error"))))
    ]
    return {
        "status": status,
        "precision_verified": status == "pass",
        "tolerance": tolerance,
        "comparisons": comparisons,
        "max_relative_error": max(finite_relative_errors, default=None),
        "max_absolute_error": max(absolute_errors, default=None),
        "max_gate_error": max(gate_errors, default=None),
        "baseline_conflict_count": len(baseline_conflicts),
        "baseline_conflicts": baseline_conflicts,
    }


def _compare_modal(result: dict[str, Any], baseline: dict[str, Any], tolerance: float) -> list[dict[str, Any]]:
    result_modes = {int(item["mode"]): float(item["frequency_hz"]) for item in result.get("modal_results", [])}
    baseline_rows = {int(item["mode"]): item for item in baseline.get("modal_results", [])}
    baseline_modes = {mode: float(item["frequency_hz"]) for mode, item in baseline_rows.items()}
    items: list[dict[str, Any]] = []
    for mode in sorted(result_modes.keys() & baseline_modes.keys()):
        baseline_row = baseline_rows.get(mode) or {}
        conflict = _baseline_conflict_for_metric(baseline_row, "frequency_hz")
        if conflict:
            items.append(
                {
                    "name": f"modal.mode_{mode}.frequency_hz",
                    "value": result_modes[mode],
                    "baseline": baseline_modes[mode],
                    "relative_error": None,
                    "absolute_error": None,
                    "gate_error": None,
                    "tolerance": tolerance,
                    "metric_type": "modal_frequency_baseline_conflict",
                    "source_ref": baseline_row.get("source_ref"),
                    "status": "baseline_conflict",
                    "baseline_conflict": conflict,
                    "message": "The baseline modal row is part of an audited source/report conflict, so this metric is excluded from the historical numerical gate only.",
                }
            )
            continue
        item = compare_scalar(f"modal.mode_{mode}.frequency_hz", result_modes[mode], baseline_modes[mode], tolerance)
        item["metric_type"] = "modal_frequency"
        items.append(item)
    return items


def _beam_value_map(result: dict[str, Any]) -> dict[tuple[str, str, str], dict[str, Any]]:
    mapped: dict[tuple[str, str, str], dict[str, Any]] = {}
    for row in result.get("beam_stress_results", []) or []:
        source_file = str(row.get("source_file") or row.get("source_ref") or "")
        key = (
            source_file,
            str(row.get("load_case") or "").upper(),
            _beam_stress_key(str(row.get("stress_type") or "")) or "",
        )
        if key[0] and key[1] and key[2]:
            mapped[key] = row
    return mapped


def _source_value_for_report_row(result_map: dict[tuple[str, str, str], dict[str, Any]], row: dict[str, Any]) -> tuple[float | None, dict[str, Any] | None]:
    key = _ratio_key(str(row.get("stress_type") or ""))
    if key in {"tension_bending", "compression_bending"}:
        return None, None
    load_case = _baseline_case_to_result_case(str(row.get("case") or ""))
    source_file = str(row.get("result_source_file") or "")
    if not key or not load_case or not source_file:
        return None, None
    result_row = result_map.get((source_file, load_case, key))
    if not result_row:
        return None, None
    return abs(float(result_row["value_mpa"])), result_row


def _baseline_eval_index(rows: list[dict[str, Any]]) -> dict[tuple[str, str, str, str, str], dict[str, Any]]:
    indexed: dict[tuple[str, str, str, str, str], dict[str, Any]] = {}
    for row in rows:
        key = _ratio_key(str(row.get("stress_type") or ""))
        case = _baseline_case_to_result_case(str(row.get("case") or ""))
        if not key or not case:
            continue
        indexed[
            (
                str(row.get("component") or ""),
                case,
                str(row.get("evaluation_kind") or ""),
                str(row.get("result_source_file") or ""),
                key,
            )
        ] = row
    return indexed


def _source_value_for_key(
    result_map: dict[tuple[str, str, str], dict[str, Any]],
    row: dict[str, Any],
    key: str,
) -> tuple[float | None, dict[str, Any] | None]:
    load_case = _baseline_case_to_result_case(str(row.get("case") or ""))
    source_file = str(row.get("result_source_file") or "")
    if not load_case or not source_file:
        return None, None
    result_row = result_map.get((source_file, load_case, key))
    if not result_row:
        return None, None
    return abs(float(result_row["value_mpa"])), result_row


def _stress_ratio_for_key(
    result_map: dict[tuple[str, str, str], dict[str, Any]],
    baseline_index: dict[tuple[str, str, str, str, str], dict[str, Any]],
    row: dict[str, Any],
    key: str,
) -> tuple[float | None, dict[str, Any] | None]:
    base_row = baseline_index.get(
        (
            str(row.get("component") or ""),
            _baseline_case_to_result_case(str(row.get("case") or "")) or "",
            str(row.get("evaluation_kind") or ""),
            str(row.get("result_source_file") or ""),
            key,
        )
    )
    if not base_row or not base_row.get("allowable_value"):
        return None, None
    value, result_row = _source_value_for_key(result_map, row, key)
    if value is None or result_row is None:
        return None, None
    coefficient = 1.0
    if row.get("evaluation_kind") == "weld_equivalent":
        coefficient = float(row.get("coefficient") or WELD_DEFAULT_COEFFICIENT)
    return (value / coefficient) / float(base_row["allowable_value"]), result_row


def _combination_value_for_report_row(
    result_map: dict[tuple[str, str, str], dict[str, Any]],
    baseline_index: dict[tuple[str, str, str, str, str], dict[str, Any]],
    row: dict[str, Any],
) -> tuple[float | None, dict[str, Any] | None]:
    key = _ratio_key(str(row.get("stress_type") or ""))
    if key == "tension_bending":
        first_key = "tension"
    elif key == "compression_bending":
        first_key = "compression"
    else:
        return None, None
    first, first_row = _stress_ratio_for_key(result_map, baseline_index, row, first_key)
    bending, bending_row = _stress_ratio_for_key(result_map, baseline_index, row, "bending")
    if first is None or bending is None:
        return None, first_row or bending_row
    return first + bending, bending_row or first_row


def _compare_ratios(result: dict[str, Any], baseline: dict[str, Any], tolerance: float) -> list[dict[str, Any]]:
    result_map = _beam_value_map(result)
    baseline_rows = baseline.get("evaluation_ratios", [])
    baseline_index = _baseline_eval_index(baseline_rows)
    items: list[dict[str, Any]] = []
    for row in baseline_rows:
        if row.get("evaluation_kind") not in {"beam_stress", "weld_equivalent"}:
            continue
        allowable = float(row.get("allowable_value") or 0)
        if allowable == 0:
            continue
        row_key = _ratio_key(str(row.get("stress_type") or ""))
        if row_key in {"tension_bending", "compression_bending"}:
            calculation_value, result_row = _combination_value_for_report_row(result_map, baseline_index, row)
            if calculation_value is None or result_row is None:
                continue
            value = calculation_value / allowable
            name = f"evaluation.{row.get('component')}.{_baseline_case_to_result_case(str(row.get('case') or '')).lower()}.{row_key}.ratio"
        else:
            base_value, result_row = _source_value_for_report_row(result_map, row)
            if base_value is None or result_row is None:
                continue
            if row.get("evaluation_kind") == "weld_equivalent":
                coefficient = float(row.get("coefficient") or WELD_DEFAULT_COEFFICIENT)
                value = (base_value / coefficient) / allowable
                name = f"evaluation.{row.get('component')}.{_baseline_case_to_result_case(str(row.get('case') or '')).lower()}.{_ratio_key(str(row.get('stress_type') or ''))}.equivalent_ratio"
            else:
                value = base_value / allowable
                name = f"evaluation.{row.get('component')}.{_baseline_case_to_result_case(str(row.get('case') or '')).lower()}.{_ratio_key(str(row.get('stress_type') or ''))}.ratio"
        baseline_value = float(row["ratio"])
        error = abs(value - baseline_value)
        conflict = _baseline_conflict_for_metric(row, "ratio")
        if conflict:
            items.append(
                {
                    "name": name,
                    "value": value,
                    "baseline": baseline_value,
                    "absolute_error": None,
                    "gate_error": None,
                    "relative_error": None,
                    "tolerance": tolerance,
                    "metric_type": "evaluation_ratio_baseline_conflict",
                    "source_ref": row.get("source_ref"),
                    "mapping_source_ref": row.get("mapping_source_ref"),
                    "result_source_file": row.get("result_source_file"),
                    "result_source_ref": result_row.get("source_ref"),
                    "status": "baseline_conflict",
                    "baseline_conflict": conflict,
                    "message": "The baseline report evaluation row is an audited source/report conflict, so this metric is excluded from the numerical precision gate.",
                }
            )
            continue
        items.append(
            {
                "name": name,
                "value": value,
                "baseline": baseline_value,
                "absolute_error": error,
                "gate_error": error,
                "relative_error": relative_error(value, baseline_value),
                "tolerance": tolerance,
                "metric_type": "evaluation_ratio",
                "source_ref": row.get("source_ref"),
                "mapping_source_ref": row.get("mapping_source_ref"),
                "result_source_file": row.get("result_source_file"),
                "result_source_ref": result_row.get("source_ref"),
                "status": "pass" if error <= tolerance else "fail",
            }
        )
    return items


def _compare_beam_calculations(result: dict[str, Any], baseline: dict[str, Any], tolerance: float) -> list[dict[str, Any]]:
    result_map = _beam_value_map(result)
    baseline_rows = baseline.get("evaluation_ratios", [])
    baseline_index = _baseline_eval_index(baseline_rows)
    items: list[dict[str, Any]] = []
    for row in baseline_rows:
        if row.get("evaluation_kind") not in {"beam_stress", "weld_equivalent"}:
            continue
        key = _ratio_key(str(row.get("stress_type") or ""))
        if not key:
            continue
        if key in {"tension_bending", "compression_bending"}:
            combo_value, result_row = _combination_value_for_report_row(result_map, baseline_index, row)
            if combo_value is None or result_row is None:
                items.append(_missing_result_item(row, key, tolerance))
                continue
            baseline_value = abs(float(row["calculation_value"]))
            name = f"combination.{row.get('component')}.{_baseline_case_to_result_case(str(row.get('case') or '')).lower()}.{key}.calculation_value"
            conflict = _baseline_conflict_for_metric(row, "calculation_value")
            if conflict:
                items.append(
                    {
                        "name": name,
                        "value": combo_value,
                        "baseline": baseline_value,
                        "relative_error": None,
                        "absolute_error": None,
                        "gate_error": None,
                        "tolerance": tolerance,
                        "metric_type": "combination_ratio_value_baseline_conflict",
                        "source_ref": row.get("source_ref"),
                        "mapping_source_ref": row.get("mapping_source_ref"),
                        "result_source_file": row.get("result_source_file"),
                        "result_source_ref": result_row.get("source_ref"),
                        "status": "baseline_conflict",
                        "baseline_conflict": conflict,
                        "message": "The baseline report evaluation row is an audited source/report conflict, so this metric is excluded from the numerical precision gate.",
                    }
                )
                continue
            item = _compare_report_scalar(name, combo_value, baseline_value, tolerance, display_decimals=row.get("calculation_decimals"))
            item["metric_type"] = "combination_ratio_value"
            item["source_ref"] = row.get("source_ref")
            item["mapping_source_ref"] = row.get("mapping_source_ref")
            item["result_source_file"] = row.get("result_source_file")
            item["result_source_ref"] = result_row.get("source_ref")
            item["result_component_scope"] = result_row.get("component_scope")
            items.append(item)
            continue
        base_value, result_row = _source_value_for_report_row(result_map, row)
        if base_value is None or result_row is None:
            items.append(_missing_result_item(row, key, tolerance))
            continue
        baseline_value = abs(float(row["calculation_value"]))
        name = f"beam.{row.get('component')}.{_baseline_case_to_result_case(str(row.get('case') or '')).lower()}.{key}.calculation_value"
        conflict = _baseline_conflict_for_metric(row, "calculation_value")
        if conflict:
            items.append(
                {
                    "name": name,
                    "value": base_value,
                    "baseline": baseline_value,
                    "relative_error": None,
                    "absolute_error": None,
                    "gate_error": None,
                    "tolerance": tolerance,
                    "metric_type": "beam_calculation_value_baseline_conflict",
                    "source_ref": row.get("source_ref"),
                    "mapping_source_ref": row.get("mapping_source_ref"),
                    "result_source_file": row.get("result_source_file"),
                    "result_source_ref": result_row.get("source_ref"),
                    "result_component_scope": result_row.get("component_scope"),
                    "status": "baseline_conflict",
                    "baseline_conflict": conflict,
                    "message": "The baseline report evaluation row is an audited source/report conflict, so this metric is excluded from the numerical precision gate.",
                }
            )
            continue
        item = _compare_report_scalar(name, base_value, baseline_value, tolerance, display_decimals=row.get("calculation_decimals"))
        item["metric_type"] = "beam_calculation_value"
        item["source_ref"] = row.get("source_ref")
        item["mapping_source_ref"] = row.get("mapping_source_ref")
        item["result_source_file"] = row.get("result_source_file")
        item["result_source_ref"] = result_row.get("source_ref")
        item["result_component_scope"] = result_row.get("component_scope")
        items.append(item)
        if row.get("evaluation_kind") == "weld_equivalent" and row.get("equivalent_stress_value") is not None:
            coefficient = float(row.get("coefficient") or WELD_DEFAULT_COEFFICIENT)
            equivalent_value = base_value / coefficient
            equivalent_baseline = abs(float(row["equivalent_stress_value"]))
            conflict = _baseline_conflict_for_metric(row, "equivalent_stress_value")
            if conflict:
                items.append(
                    {
                        "name": f"weld.{row.get('component')}.{_baseline_case_to_result_case(str(row.get('case') or '')).lower()}.{key}.equivalent_stress_value",
                        "value": equivalent_value,
                        "baseline": equivalent_baseline,
                        "relative_error": None,
                        "absolute_error": None,
                        "gate_error": None,
                        "tolerance": tolerance,
                        "metric_type": "weld_equivalent_stress_value_baseline_conflict",
                        "source_ref": row.get("source_ref"),
                        "mapping_source_ref": row.get("coefficient_source_ref") or row.get("mapping_source_ref"),
                        "result_source_file": row.get("result_source_file"),
                        "result_source_ref": result_row.get("source_ref"),
                        "status": "baseline_conflict",
                        "baseline_conflict": conflict,
                        "message": "The baseline report evaluation row is an audited source/report conflict, so this metric is excluded from the numerical precision gate.",
                    }
                )
                continue
            equivalent_item = _compare_report_scalar(
                f"weld.{row.get('component')}.{_baseline_case_to_result_case(str(row.get('case') or '')).lower()}.{key}.equivalent_stress_value",
                equivalent_value,
                equivalent_baseline,
                tolerance,
                display_decimals=row.get("equivalent_stress_decimals"),
            )
            equivalent_item["metric_type"] = "weld_equivalent_stress_value"
            equivalent_item["source_ref"] = row.get("source_ref")
            equivalent_item["mapping_source_ref"] = row.get("coefficient_source_ref") or row.get("mapping_source_ref")
            equivalent_item["result_source_file"] = row.get("result_source_file")
            equivalent_item["result_source_ref"] = result_row.get("source_ref")
            items.append(equivalent_item)
    return items


def _missing_result_item(row: dict[str, Any], key: str, tolerance: float) -> dict[str, Any]:
    return {
        "name": f"mapping.{row.get('component')}.{_baseline_case_to_result_case(str(row.get('case') or ''))}.{key}",
        "value": None,
        "baseline": row.get("calculation_value"),
        "relative_error": None,
        "absolute_error": None,
        "tolerance": tolerance,
        "metric_type": "missing_result_mapping",
        "source_ref": row.get("source_ref"),
        "mapping_source_ref": row.get("mapping_source_ref"),
        "result_source_file": row.get("result_source_file"),
        "status": "fail",
        "message": "No parsed result row exists for the report-declared source file/load case/stress type.",
    }


def _compare_report_scalar(
    name: str,
    value: float,
    baseline: float,
    tolerance: float,
    *,
    display_decimals: int | None = None,
) -> dict[str, Any]:
    absolute_error = abs(value - baseline)
    display_rounding_match = False
    display_resolution_match = False
    if display_decimals is not None and display_decimals >= 0:
        display_rounding_match = round(value, display_decimals) == round(baseline, display_decimals)
        display_resolution = 10 ** (-display_decimals)
        display_resolution_match = absolute_error <= display_resolution + 1e-12
    if baseline == 0:
        status = "pass" if absolute_error <= tolerance or display_rounding_match or display_resolution_match else "fail"
        rel = 0.0 if value == 0 else float("inf")
    else:
        rel = absolute_error / abs(baseline)
        status = "pass" if rel <= tolerance or display_rounding_match or display_resolution_match else "fail"
    return {
        "name": name,
        "value": value,
        "baseline": baseline,
        "relative_error": rel,
        "absolute_error": absolute_error,
        "gate_error": 0.0 if display_rounding_match or display_resolution_match else (absolute_error if baseline == 0 else rel),
        "tolerance": tolerance,
        "status": status,
        "display_decimals": display_decimals,
        "display_rounding_match": display_rounding_match,
        "display_resolution_match": display_resolution_match,
    }


def _compare_loads(result: dict[str, Any], baseline: dict[str, Any], tolerance: float) -> list[dict[str, Any]]:
    if not baseline.get("loads"):
        return []
    items: list[dict[str, Any]] = []
    foundation_baseline = [row for row in baseline["loads"] if row.get("load_kind") == "foundation"]
    items.extend(
        _compare_load_collection(
            result.get("foundation_loads") or [],
            foundation_baseline,
            tolerance,
            metric_type="foundation_load",
            name_prefix="foundation_loads",
        )
    )

    root_baseline = [row for row in baseline["loads"] if row.get("load_kind") == "cantilever_root_load"]
    items.extend(
        _compare_load_collection(
            result.get("weld_force_results") or [],
            root_baseline,
            tolerance,
            metric_type="cantilever_root_load",
            name_prefix="cantilever_root_loads",
        )
    )

    legacy_connection_baseline = [
        row
        for row in baseline["loads"]
        if row.get("load_kind") == "tray_arm_connection" and row.get("result_source_file") != "LS-FORCE-NODES.LIS"
    ]
    items.extend(
        _compare_load_collection(
            result.get("bolt_force_results") or [],
            legacy_connection_baseline,
            tolerance,
            metric_type="tray_arm_connection_load",
            name_prefix="tray_arm_connection_loads",
        )
    )

    interface_baseline = [
        row
        for row in baseline["loads"]
        if row.get("load_kind") == "tray_arm_connection" and row.get("result_source_file") == "LS-FORCE-NODES.LIS"
    ]
    for index, baseline_row in enumerate(interface_baseline):
        interface_values = _interface_bolt_load_from_topology(result, baseline_row)
        if interface_values is None:
            if baseline_row.get("baseline_conflicts"):
                for key, baseline_value in (baseline_row.get("values") or {}).items():
                    conflict = _baseline_conflict_for_metric(baseline_row, key)
                    if not conflict:
                        continue
                    items.append(
                        {
                            "name": f"tray_arm_connection_interface[{index}].{key}",
                            "value": None,
                            "baseline": baseline_value,
                            "relative_error": None,
                            "absolute_error": None,
                            "gate_error": None,
                            "tolerance": tolerance,
                            "metric_type": "tray_arm_connection_interface_baseline_conflict",
                            "source_ref": baseline_row.get("source_ref"),
                            "mapping_source_ref": baseline_row.get("mapping_source_ref"),
                            "result_source_file": "LS-FORCE-NODES.LIS",
                            "status": "baseline_conflict",
                            "baseline_conflict": conflict,
                            "message": "The baseline report row is an audited source/report conflict, so this metric is excluded from the numerical precision gate.",
                        }
                    )
                continue
            items.append(
                {
                    "name": f"tray_arm_connection_interface[{index}]",
                    "value": None,
                    "baseline": baseline_row.get("values"),
                    "relative_error": None,
                    "absolute_error": None,
                    "tolerance": tolerance,
                    "metric_type": "tray_arm_connection_interface_missing",
                    "source_ref": baseline_row.get("source_ref"),
                    "mapping_source_ref": baseline_row.get("mapping_source_ref"),
                    "result_source_file": "LS-FORCE-NODES.LIS",
                    "status": "fail",
                    "message": "Report table is a tray-arm CP-interface bolt-load table, but no topology export rows are available.",
                }
            )
            continue
        for key, baseline_value in (baseline_row.get("values") or {}).items():
            conflict = _baseline_conflict_for_metric(baseline_row, key)
            if conflict:
                items.append(
                    {
                        "name": f"tray_arm_connection_interface[{index}].{key}",
                        "value": abs(float(interface_values[key])) if key in interface_values else None,
                        "baseline": abs(float(baseline_value)),
                        "relative_error": None,
                        "absolute_error": None,
                        "gate_error": None,
                        "tolerance": tolerance,
                        "metric_type": "tray_arm_connection_interface_baseline_conflict",
                        "source_ref": baseline_row.get("source_ref"),
                        "mapping_source_ref": baseline_row.get("mapping_source_ref"),
                        "result_source_file": "LS-FORCE-NODES.LIS",
                        "table_caption": baseline_row.get("table_caption"),
                        "status": "baseline_conflict",
                        "baseline_conflict": conflict,
                        "message": "The baseline report row is an audited source/report conflict, so this metric is excluded from the numerical precision gate.",
                    }
                )
                continue
            if key not in interface_values:
                continue
            item = compare_scalar(
                f"tray_arm_connection_interface[{index}].{key}",
                abs(float(interface_values[key])),
                abs(float(baseline_value)),
                tolerance,
            )
            item["gate_error"] = item.get("relative_error")
            item["metric_type"] = "tray_arm_connection_interface"
            item["source_ref"] = baseline_row.get("source_ref")
            item["mapping_source_ref"] = baseline_row.get("mapping_source_ref")
            item["result_source_file"] = "LS-FORCE-NODES.LIS"
            item["table_caption"] = baseline_row.get("table_caption")
            item["mapping_policy"] = (
                "Use the model-declared tray-arm CP interface keypoint family (*2) from LS-FORCE-NODES.LIS. "
                "Do not substitute standard LS-FORCE.LIS or choose rows by numeric closeness."
            )
            items.append(item)

    for index, baseline_row in enumerate(row for row in baseline["loads"] if row.get("load_kind") == "tray_arm_connection_derived_bolt_load"):
        derived_values = _derived_bolt_load_from_nodes(result, baseline_row)
        if derived_values is None:
            if baseline_row.get("baseline_conflicts"):
                for key, baseline_value in (baseline_row.get("values") or {}).items():
                    conflict = _baseline_conflict_for_metric(baseline_row, key)
                    if not conflict:
                        continue
                    items.append(
                        {
                            "name": f"derived_bolt_loads[{index}].{key}",
                            "value": None,
                            "baseline": baseline_value,
                            "relative_error": None,
                            "absolute_error": None,
                            "gate_error": None,
                            "tolerance": tolerance,
                            "metric_type": "derived_bolt_load_baseline_conflict",
                            "source_ref": baseline_row.get("source_ref"),
                            "mapping_source_ref": baseline_row.get("mapping_source_ref"),
                            "result_source_file": "LS-FORCE-NODES.LIS",
                            "status": "baseline_conflict",
                            "baseline_conflict": conflict,
                            "message": "The baseline report row is an audited source/report conflict, so this metric is excluded from the numerical precision gate.",
                        }
                    )
                continue
            items.append(
                {
                    "name": f"derived_bolt_loads[{index}]",
                    "value": None,
                    "baseline": baseline_row.get("values"),
                    "relative_error": None,
                    "absolute_error": None,
                    "tolerance": tolerance,
                    "metric_type": "derived_bolt_load_missing",
                    "source_ref": baseline_row.get("source_ref"),
                    "mapping_source_ref": baseline_row.get("mapping_source_ref"),
                    "result_source_file": "LS-FORCE-NODES.LIS",
                    "status": "fail",
                    "message": "Report table is a secondary tray-arm bolt-load table, but no topology-mapped connection-node export is available.",
                }
            )
            continue
        for key, baseline_value in (baseline_row.get("values") or {}).items():
            conflict = _baseline_conflict_for_metric(baseline_row, key)
            if conflict:
                items.append(
                    {
                        "name": f"derived_bolt_loads[{index}].{key}",
                        "value": abs(float(derived_values[key])) if key in derived_values else None,
                        "baseline": abs(float(baseline_value)),
                        "relative_error": None,
                        "absolute_error": None,
                        "gate_error": None,
                        "tolerance": tolerance,
                        "metric_type": "derived_bolt_load_baseline_conflict",
                        "source_ref": baseline_row.get("source_ref"),
                        "mapping_source_ref": baseline_row.get("mapping_source_ref"),
                        "result_source_file": "LS-FORCE-NODES.LIS",
                        "table_caption": baseline_row.get("table_caption"),
                        "status": "baseline_conflict",
                        "baseline_conflict": conflict,
                        "message": "The baseline report row is an audited source/report conflict, so this metric is excluded from the numerical precision gate.",
                    }
                )
                continue
            if key not in derived_values:
                continue
            item = compare_scalar(
                f"derived_bolt_loads[{index}].{key}",
                abs(float(derived_values[key])),
                abs(float(baseline_value)),
                tolerance,
            )
            item["gate_error"] = item.get("relative_error")
            item["metric_type"] = "derived_bolt_load"
            item["source_ref"] = baseline_row.get("source_ref")
            item["mapping_source_ref"] = baseline_row.get("mapping_source_ref")
            item["result_source_file"] = "LS-FORCE-NODES.LIS"
            item["table_caption"] = baseline_row.get("table_caption")
            items.append(item)
    return items


def _node_metric_value(row: dict[str, Any], key: str) -> float:
    value = row.get(key)
    if isinstance(value, dict):
        return float(value.get("value") or 0.0)
    return float(value or 0.0)


def _has_nonzero_node_force(row: dict[str, Any] | None) -> bool:
    if not row:
        return False
    return any(abs(_node_metric_value(row, key)) > 1e-9 for key in ("fx", "fy", "fz", "mx", "my", "mz"))


def _interface_bolt_load_from_topology(result: dict[str, Any], baseline_row: dict[str, Any]) -> dict[str, float] | None:
    load_case = _baseline_case_to_result_case(str(baseline_row.get("case") or ""))
    if not load_case:
        return None
    group = [
        row
        for row in result.get("connection_node_force_results", []) or []
        if str(row.get("load_case") or "").upper() == load_case
        and _has_nonzero_node_force(row)
        and int(row.get("keypoint") or 0) % 10 == 2
        and (500 <= int(row.get("keypoint") or 0) < 800 or 1500 <= int(row.get("keypoint") or 0) < 1800)
    ]
    if not group:
        return None
    return {
        key: max((_node_metric_value(row, key) for row in group), key=abs)
        for key in ("fx", "fy", "fz", "mx", "my", "mz")
    }


def _derived_bolt_load_from_nodes(result: dict[str, Any], baseline_row: dict[str, Any]) -> dict[str, float] | None:
    load_case = _baseline_case_to_result_case(str(baseline_row.get("case") or ""))
    if not load_case:
        return None
    rows = {
        (int(row.get("keypoint") or 0), str(row.get("load_case") or "").upper()): row
        for row in result.get("connection_node_force_results", []) or []
    }
    # Secondary perpendicular tray models use keypoint 109/209/309; the report
    # table is one representative connection point, not the max over the main
    # tray-arm connection set.
    for keypoint in (109, 209, 309):
        row = rows.get((keypoint, load_case))
        if _has_nonzero_node_force(row):
            return {key: _node_metric_value(row, key) for key in ("fx", "fy", "fz", "mx", "my", "mz")}

    # Single-side mixed-width models place the secondary tray at layer suffix 69.
    # The reviewed reports envelope each component over the secondary connection
    # group. This preserves the report-table logic without matching by numeric
    # closeness to the baseline row.
    group = [
        row
        for keypoint in (569, 669, 769, 1509, 1519, 1529, 1559, 1569, 1579, 1669, 1769)
        if _has_nonzero_node_force(row := rows.get((keypoint, load_case)))
    ]
    if group:
        return {
            key: max((_node_metric_value(row, key) for row in group), key=abs)
            for key in ("fx", "fy", "fz", "mx", "my", "mz")
        }
    return None


def _compare_load_collection(
    result_rows: list[dict[str, Any]],
    baseline_rows: list[dict[str, Any]],
    tolerance: float,
    *,
    metric_type: str,
    name_prefix: str,
) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    result_by_case: dict[str, dict[str, Any]] = {}
    for result_row in result_rows:
        result_case = str(result_row.get("load_case") or result_row.get("case") or "").upper()
        if result_case and result_case not in result_by_case:
            result_by_case[result_case] = result_row

    for index, baseline_row in enumerate(baseline_rows):
        expected_case = _baseline_case_to_result_case(str(baseline_row.get("case") or ""))
        result_row = result_by_case.get(expected_case or "")
        if result_row is None and index < len(result_rows):
            result_row = result_rows[index]
        if result_row is None:
            if baseline_row.get("baseline_conflicts"):
                for key, baseline_value in (baseline_row.get("values") or {}).items():
                    conflict = _baseline_conflict_for_metric(baseline_row, key)
                    if not conflict:
                        continue
                    items.append(
                        {
                            "name": f"{name_prefix}[{index}].{key}",
                            "value": None,
                            "baseline": baseline_value,
                            "relative_error": None,
                            "absolute_error": None,
                            "gate_error": None,
                            "tolerance": tolerance,
                            "metric_type": f"{metric_type}_baseline_conflict",
                            "source_ref": baseline_row.get("source_ref"),
                            "mapping_source_ref": baseline_row.get("mapping_source_ref"),
                            "result_source_file": baseline_row.get("result_source_file"),
                            "table_caption": baseline_row.get("table_caption"),
                            "status": "baseline_conflict",
                            "baseline_conflict": conflict,
                            "message": "The baseline report row is an audited source/report conflict, so this metric is excluded from the numerical precision gate.",
                        }
                    )
                continue
            items.append(
                {
                    "name": f"{name_prefix}[{index}]",
                    "value": None,
                    "baseline": baseline_row.get("values"),
                    "relative_error": None,
                    "absolute_error": None,
                    "tolerance": tolerance,
                    "metric_type": f"{metric_type}_missing",
                    "source_ref": baseline_row.get("source_ref"),
                    "mapping_source_ref": baseline_row.get("mapping_source_ref"),
                    "result_source_file": baseline_row.get("result_source_file"),
                    "status": "fail",
                    "message": "No parsed result row exists for this report-declared load table row.",
                }
            )
            continue
        baseline_row = baseline_rows[index]
        values = result_row.get("values") or result_row
        for key, metric in (baseline_row.get("values") or {}).items():
            conflict = _baseline_conflict_for_metric(baseline_row, key)
            if conflict:
                result_metric = values.get(key) if isinstance(values, dict) else None
                value = None
                if result_metric is not None:
                    value = float(result_metric["value"] if isinstance(result_metric, dict) else result_metric)
                    unit = str(result_metric.get("unit") if isinstance(result_metric, dict) else "")
                    value = _normalise_load_for_report_unit(value, unit)
                items.append(
                    {
                        "name": f"{name_prefix}[{index}].{key}",
                        "value": abs(value) if value is not None else None,
                        "baseline": abs(float(metric)),
                        "relative_error": None,
                        "absolute_error": None,
                        "gate_error": None,
                        "tolerance": tolerance,
                        "metric_type": f"{metric_type}_baseline_conflict",
                        "source_ref": baseline_row.get("source_ref"),
                        "mapping_source_ref": baseline_row.get("mapping_source_ref"),
                        "result_source_file": baseline_row.get("result_source_file"),
                        "table_caption": baseline_row.get("table_caption"),
                        "result_source_ref": result_row.get("source_ref") if isinstance(result_row, dict) else None,
                        "status": "baseline_conflict",
                        "baseline_conflict": conflict,
                        "message": "The baseline report row is an audited source/report conflict, so this metric is excluded from the numerical precision gate.",
                    }
                )
                continue
            if key not in values:
                continue
            result_metric = values[key]
            value = float(result_metric["value"] if isinstance(result_metric, dict) else result_metric)
            unit = str(result_metric.get("unit") if isinstance(result_metric, dict) else "")
            value = _normalise_load_for_report_unit(value, unit)
            baseline_value = float(metric)
            item = compare_scalar(f"{name_prefix}[{index}].{key}", abs(value), abs(baseline_value), tolerance)
            item["gate_error"] = item.get("relative_error")
            item["metric_type"] = metric_type
            item["source_ref"] = baseline_row.get("source_ref")
            item["mapping_source_ref"] = baseline_row.get("mapping_source_ref")
            item["result_source_file"] = baseline_row.get("result_source_file")
            item["table_caption"] = baseline_row.get("table_caption")
            result_source = result_row.get("source_ref") if isinstance(result_row, dict) else None
            if result_source:
                item["result_source_ref"] = result_source
            items.append(item)
    return items


def _mapping_blockers(result: dict[str, Any], baseline: dict[str, Any]) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    source_scopes = {
        str(row.get("source_file") or row.get("source_ref") or ""): row.get("component_scope")
        for row in result.get("beam_stress_results", []) or []
    }
    for row in baseline.get("evaluation_ratios", []):
        if (
            row.get("component") == "square_support"
            and row.get("result_source_file") == "SQUAREBEAMSTRESS.LIS"
            and "SQUAREBEAMSTRESS.LIS" not in source_scopes
        ):
            items.append(
                {
                    "name": "mapping.square_support.SQUAREBEAMSTRESS.missing",
                    "value": None,
                    "baseline": "square_support_section_1",
                    "relative_error": None,
                    "tolerance": STRICT_TOLERANCE,
                    "status": "fail",
                    "metric_type": "mapping_scope",
                    "source_ref": row.get("source_ref"),
                    "mapping_source_ref": row.get("mapping_source_ref"),
                    "message": "Report caption says square support, but no section-1-only SQUAREBEAMSTRESS.LIS result is present. The all-section MAXBEAMSTRESS.LIS must not be substituted by numeric closeness.",
                }
            )
            break
        if (
            row.get("component") == "square_support"
            and row.get("result_source_file") == "SQUAREBEAMSTRESS.LIS"
            and source_scopes.get("SQUAREBEAMSTRESS.LIS") != "beam_type_1_section_1_only"
        ):
            items.append(
                {
                    "name": "mapping.square_support.SQUAREBEAMSTRESS.scope",
                    "value": source_scopes.get("SQUAREBEAMSTRESS.LIS"),
                    "baseline": "beam_type_1_section_1_only",
                    "relative_error": None,
                    "tolerance": STRICT_TOLERANCE,
                    "status": "fail",
                    "metric_type": "mapping_scope",
                    "source_ref": row.get("source_ref"),
                    "mapping_source_ref": row.get("mapping_source_ref"),
                    "message": "SQUAREBEAMSTRESS.LIS exists but was not generated from a verified section-1-only selector.",
                }
            )
            break
    return items


def _normalise_load_for_report_unit(value: float, unit: str) -> float:
    unit_text = unit.replace(" ", "").lower()
    if unit_text == "kn":
        return value * 1000.0
    if unit_text in {"kn*m", "kn.m", "kn-m"}:
        return value * 1000.0
    return value


def _ratio_key(text: str) -> str | None:
    if "拉弯" in text or ("拉" in text and "弯" in text and "组合" in text):
        return "tension_bending"
    if "压弯" in text or ("压" in text and "弯" in text and "组合" in text):
        return "compression_bending"
    if "拉伸" in text or "拉应力" in text:
        return "tension"
    if "压缩" in text:
        return "compression"
    if "弯曲" in text:
        return "bending"
    if "剪切" in text or "剪应力" in text:
        return "shear"
    return None


def _beam_stress_key(text: str) -> str | None:
    upper = text.upper()
    if "TENSION" in upper or "SDIR_TEN" in upper:
        return "tension"
    if "COMPRESSION" in upper or "SDIR_COM" in upper:
        return "compression"
    if "BEND" in upper:
        return "bending"
    if "SHEAR" in upper:
        return "shear"
    return None


def _baseline_case_to_result_case(text: str) -> str | None:
    upper = str(text or "").upper()
    if upper in {"DW", "SL-1", "SL-2", "NORMAL", "UPSET", "FAULTED"}:
        return upper
    if "事故" in text:
        return "FAULTED"
    if "异常" in text:
        return "UPSET"
    if "正常" in text:
        return "NORMAL"
    return None


def write_report_baseline_comparison(
    job_dir: Path | str,
    report_path: Path | str,
    *,
    tolerance: float = STRICT_TOLERANCE,
) -> dict[str, Any]:
    job_dir = Path(job_dir)
    result = json.loads((job_dir / "result.json").read_text(encoding="utf-8"))
    baseline = extract_report_baseline(report_path)
    comparison = compare_report_baseline(result, baseline, tolerance=tolerance)
    payload = {
        **comparison,
        "baseline_report": str(report_path),
        "baseline_file": "baseline_from_report.json",
    }
    (job_dir / "baseline_from_report.json").write_text(
        json.dumps(baseline, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    (job_dir / "baseline_comparison.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    lines = [
        "# Precision Acceptance Report",
        "",
        f"Status: {payload['status']}",
        f"Precision verified: {payload['precision_verified']}",
        f"Tolerance: {tolerance:.2%}",
        f"Baseline report: {report_path}",
        "",
    ]
    if payload["status"] == "blocked":
        lines.append(payload.get("message", "No comparable fields were found."))
    for item in payload.get("comparisons", []):
        error = item.get("absolute_error", item.get("relative_error"))
        lines.append(
            f"- {item['name']}: {item['status']} value={item['value']} baseline={item['baseline']} error={error}"
        )
    text = "\n".join(lines) + "\n"
    (job_dir / "baseline_comparison.md").write_text(text, encoding="utf-8")
    (job_dir / "precision_acceptance_report.md").write_text(text, encoding="utf-8")
    return payload
