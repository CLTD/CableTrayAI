from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Inches

from core.report.report_audit import audit_report
from core.report.reference_compare import compare_report_structure


def _add_table(document: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = "" if value is None else str(value)


def _ratio_text(value: float | None) -> str:
    return "" if value is None else f"{value:.3f}"


def _metric_text(metric: dict) -> str:
    return f"{metric.get('value')} {metric.get('unit') or ''}".strip()


def _eval_rows(evaluation: list[dict], prefix: str) -> list[dict]:
    return [item for item in evaluation if item["check_id"].startswith(prefix) or f"_{prefix}_" in item["check_id"]]


def _conclusion(evaluation: list[dict]) -> str:
    if any(item.get("pass_fail") == "不满足" for item in evaluation):
        return "不满足"
    if any(item.get("formula_status") == "unconfirmed_todo" or item.get("pass_fail") == "待确认" for item in evaluation):
        return "待确认"
    return "满足"


def build_report(job_dir: Path | str) -> dict:
    job_dir = Path(job_dir)
    result = json.loads((job_dir / "result.json").read_text(encoding="utf-8"))
    input_payload = json.loads((job_dir / "input.json").read_text(encoding="utf-8"))
    evaluation = result.get("evaluation_summary", [])
    figures = json.loads((job_dir / "figures_manifest.json").read_text(encoding="utf-8"))

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    project = result["project"]
    document.add_heading("Cable Tray Seismic Analysis Report", level=1)
    document.add_paragraph(f"Project code: {project['project_code']}")
    document.add_paragraph(f"Building: {project['building']}")
    document.add_paragraph(f"Area: {project['area']}")
    document.add_paragraph(f"Elevation: {project['elevation']}")

    document.add_heading("1 概述", level=1)
    document.add_paragraph("本报告由 CableTrayAI 根据 input.json、result.json 和 figures_manifest.json 自动生成。")

    document.add_heading("2 结构说明", level=1)
    document.add_heading("2.1 材料参数", level=2)
    _add_table(
        document,
        ["Material", "E (Pa)", "Poisson", "Density", "Normal allowable", "Shear allowable", "Source"],
        [
            [
                material["name"],
                material["elastic_modulus_pa"],
                material["poisson_ratio"],
                material["density_kg_m3"],
                material.get("allowable_normal_mpa"),
                material.get("allowable_shear_mpa"),
                material.get("source_ref"),
            ]
            for material in input_payload.get("materials", [])
        ],
    )

    document.add_heading("2.2 结构设计参数", level=2)
    support = input_payload.get("support", {})
    _add_table(
        document,
        ["Parameter", "Value"],
        [[key, value] for key, value in support.items() if key != "source_ref"],
    )

    document.add_heading("3 工况", level=1)
    document.add_heading("3.1 载荷工况", level=2)
    _add_table(
        document,
        ["Name", "Category", "Level", "Directions", "Source"],
        [
            [
                row.get("name"),
                row.get("category"),
                row.get("spectrum_level"),
                ",".join(row.get("directions", [])),
                row.get("source_ref"),
            ]
            for row in input_payload.get("load_cases", [])
        ],
    )

    document.add_heading("3.2 工况组合", level=2)
    document.add_paragraph("工况组合来自 result.json 与 PIP/LIS 解析结果。")

    document.add_heading("4 规范要求", level=1)
    document.add_paragraph("规范条文和公式来源通过 source_ref 追溯；未确认公式不作为最终通过结论。")

    document.add_heading("5 计算方法、程序和计算模型", level=1)
    document.add_paragraph("计算模型、求解和后处理文件由 APDL 模板渲染生成。")

    document.add_heading("6 结果评定", level=1)
    document.add_heading("6.1 支架的评定", level=2)
    _add_evaluation_table(document, _eval_rows(evaluation, "support"))

    document.add_heading("6.2 焊缝的评定", level=2)
    _add_evaluation_table(document, _eval_rows(evaluation, "weld"))

    document.add_heading("6.3 基础载荷", level=2)
    _add_table(
        document,
        ["Case", "Node", "Fx", "Fy", "Fz", "Mx", "My", "Mz"],
        [
            [
                row["load_case"],
                row["node"],
                _metric_text(row["fx"]),
                _metric_text(row["fy"]),
                _metric_text(row["fz"]),
                _metric_text(row["mx"]),
                _metric_text(row["my"]),
                _metric_text(row["mz"]),
            ]
            for row in result["foundation_loads"]
        ],
    )

    document.add_heading("6.4 螺栓的评定", level=2)
    document.add_heading("螺栓载荷表", level=3)
    _add_table(
        document,
        ["Name", "Case", "Tension", "Shear", "Allow tension", "Allow shear", "Source"],
        [
            [
                row["name"],
                row["load_case"],
                _metric_text(row["values"]["tension_mpa"]),
                _metric_text(row["values"]["shear_mpa"]),
                _metric_text(row["values"]["allowable_tension_mpa"]),
                _metric_text(row["values"]["allowable_shear_mpa"]),
                row["source_ref"],
            ]
            for row in result.get("bolt_force_results", [])
        ],
    )

    document.add_heading("螺栓评定表", level=3)
    _add_evaluation_table(document, _eval_rows(evaluation, "bolt"))

    document.add_heading("7 结论", level=1)
    conclusion = _conclusion(evaluation)
    max_ratio = max([float(item["ratio"]) for item in evaluation if item.get("ratio") is not None], default=None)
    document.add_paragraph(f"Evaluation conclusion: {conclusion}")
    document.add_paragraph(f"Maximum ratio: {_ratio_text(max_ratio)}")

    document.add_heading("参考资料", level=1)
    document.add_paragraph("参考资料清单由 source_inventory.json 和 source_ref 记录维护。")

    document.add_heading("附录A 模态分析结果", level=1)
    document.add_heading("模态频率表", level=2)
    _add_table(
        document,
        ["Mode", "Frequency (Hz)", "Period (s)", "Source"],
        [
            [
                row["mode"],
                f"{row['frequency_hz']:.3f}",
                f"{row['period_s']:.4f}",
                row["source_ref"],
            ]
            for row in result["modal_results"]
        ],
    )

    document.add_heading("模态图", level=2)
    _add_figures(document, job_dir, [figure for figure in figures if figure.get("figure_type") == "modal" or figure.get("category") == "modal"])

    document.add_heading("附录B 支架应力图", level=1)
    _add_figures(document, job_dir, [figure for figure in figures if figure.get("figure_type") == "stress" or figure.get("category") == "stress"])

    document.add_heading("附录C 焊缝评定原理", level=1)
    document.add_paragraph("焊缝评定公式由 formula_registry 和 source_ref 追溯；未确认公式保持待确认。")

    output_path = job_dir / "report.docx"
    document.save(output_path)
    audit = audit_report(job_dir)
    audit["reference_comparison"] = compare_report_structure(
        output_path,
        output_path=job_dir / "report_reference_comparison.json",
    )
    (job_dir / "report_audit.json").write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    return audit


def _add_evaluation_table(document: Document, rows: list[dict]) -> None:
    _add_table(
        document,
        ["Check", "Calc", "Allowable", "Ratio", "Result", "Source"],
        [
            [
                item["category"],
                item["calculation_value"],
                item["allowable_value"],
                _ratio_text(item["ratio"]),
                item["pass_fail"],
                item["source_ref"],
            ]
            for item in rows
        ],
    )


def _add_figures(document: Document, job_dir: Path, figures: list[dict]) -> None:
    for figure in figures:
        figure_path = job_dir / figure.get("target_file", figure["path"])
        document.add_paragraph(figure["caption"])
        if figure_path.exists():
            try:
                document.add_picture(str(figure_path), width=Inches(4.8))
            except Exception:
                document.add_paragraph(f"Figure file: {figure_path.name}")
