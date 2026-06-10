from __future__ import annotations

import json
import re
import shutil
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor

from core.validation.analysis_scope import classify_scope_from_input


TEMPLATE_DIR = Path("templates/report")
STEEL_TEMPLATE = TEMPLATE_DIR / "steel_platform_report_template.docx"
NON_STEEL_TEMPLATE = TEMPLATE_DIR / "non_steel_platform_report_template.docx"


STRESS_LABELS = [
    ("SDIR_TENSION", "拉伸应力", "support_tension"),
    ("SDIR_COMPRESSION", "压缩应力", "support_compression"),
    ("SBEND", "弯曲应力", "support_bending"),
    ("SHEAR", "剪切应力", "support_shear"),
]

TITLE_RED = RGBColor(255, 0, 0)
WELD_EQUIVALENT_COEFFICIENT = 0.526


def _read_json(path: Path, default: Any) -> Any:
    if not path.exists():
        return default
    return json.loads(path.read_text(encoding="utf-8-sig"))


def _safe_name(value: str) -> str:
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", str(value or "")).strip(" ._")
    return name or "CableTrayAI_report"


def _is_static_or_steel(input_payload: dict[str, Any]) -> bool:
    metadata = input_payload.get("metadata") or {}
    method = str(metadata.get("analysis_method") or "").lower()
    raw_row = metadata.get("raw_intake_row") or {}
    combined = " ".join(str(value) for value in [metadata, input_payload.get("project") or {}, raw_row])
    return method == "static" or "钢平台" in combined


def choose_report_template(input_payload: dict[str, Any], template_dir: Path | str = TEMPLATE_DIR) -> dict[str, str]:
    template_dir = Path(template_dir)
    scope = classify_scope_from_input(input_payload)
    if scope.get("analysis_method") == "static" or _is_static_or_steel(input_payload):
        template = template_dir / STEEL_TEMPLATE.name
        mode = "steel_platform"
    else:
        template = template_dir / NON_STEEL_TEMPLATE.name
        mode = "non_steel_platform"
    return {"mode": mode, "template": str(template)}


def _report_id(job_dir: Path, input_payload: dict[str, Any]) -> str:
    metadata = input_payload.get("metadata") or {}
    for key in ("report_number", "calculation_batch", "intake_order_id", "provisional_intake_id"):
        if metadata.get(key):
            return _safe_name(str(metadata[key]))
    return _safe_name(job_dir.name)


def _set_cell_text(cell, value: Any) -> None:
    text = "" if value is None else str(value)
    if not cell.paragraphs:
        cell.add_paragraph()
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    for extra in cell.paragraphs[1:]:
        for run in extra.runs:
            run.text = ""


def _fmt_num(value: Any, digits: int = 2) -> str:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return "" if value is None else str(value)
    if abs(number) < 1e-12:
        return "0"
    text = f"{number:.{digits}f}"
    if float(text) == 0.0:
        return f"{number:.6g}"
    return text


def _fmt_ratio(value: Any) -> str:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return "-"
    if abs(number) < 1e-12:
        return "0"
    text = f"{number:.2f}"
    if float(text) == 0.0:
        return f"{number:.6g}"
    return text


def _fmt_ratio_precise(value: Any, digits: int = 3) -> str:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return "-"
    if abs(number) < 1e-12:
        return "0"
    text = f"{number:.{digits}f}".rstrip("0").rstrip(".")
    return text or "0"


def _metric_value(metric: Any) -> Any:
    if isinstance(metric, dict):
        return metric.get("value", metric.get("normalized_value"))
    return metric


def _float_list(values: Any) -> list[float]:
    if values is None:
        return []
    if not isinstance(values, (list, tuple)):
        values = [values]
    parsed: list[float] = []
    for item in values:
        try:
            parsed.append(float(item))
        except (TypeError, ValueError):
            continue
    return parsed


def _static_report_spectrum_elevations(input_payload: dict[str, Any]) -> list[float]:
    metadata = input_payload.get("metadata") or {}
    if str(metadata.get("analysis_method") or "").lower() != "static":
        return []
    source = metadata.get("static_acceleration_source") or {}
    selected = _float_list(source.get("elevations"))
    selected.extend(_float_list(source.get("selected_elevations")))
    selected.extend(_float_list(source.get("selected_elevation")))
    selected.extend(_float_list(source.get("elevation")))
    if selected:
        return sorted(set(selected))
    return sorted(set(_float_list(metadata.get("static_elevation_candidates"))))


def _mark_paragraph_red(paragraph) -> bool:
    if not paragraph.runs:
        paragraph.add_run(paragraph.text)
    for run in paragraph.runs:
        run.font.color.rgb = TITLE_RED
    return True


def _mark_first_title_red(document: Document, *prefixes: str) -> str | None:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if any(text.startswith(prefix) for prefix in prefixes if prefix):
            _mark_paragraph_red(paragraph)
            return text
    return None


def _mark_titles_red(document: Document, prefixes: set[str]) -> list[str]:
    marked: list[str] = []
    for paragraph in document.paragraphs:
        style_name = (paragraph.style.name or "").lower()
        if style_name.startswith("toc"):
            continue
        text = paragraph.text.strip()
        if not text:
            continue
        if any(text.startswith(prefix) for prefix in prefixes if prefix):
            _mark_paragraph_red(paragraph)
            marked.append(text)
    return marked


def _replacement_title_prefixes(replacements: list[dict[str, Any]]) -> set[str]:
    prefixes: set[str] = set()
    target_to_prefixes = {
        "表3-1": {"表3-1"},
        "支架应力评定": {"表6-1"},
        "support stress evaluation": {"表6-1"},
        "square support stress evaluation": {"表6-2"},
        "cantilever/root weld stress evaluation": {"表6-2"},
        "托臂根部焊缝评定结果": {"表6-2", "表6-3"},
        "焊缝评定结果": {"表6-2", "表6-3"},
        "托臂根部所受载荷": {"表6-2"},
        "支架基础载荷": {"表6-4"},
        "支架连接螺栓载荷": {"表6-5"},
        "支架螺栓应力评定表": {"表6-6"},
        "模态频率表": {"表A-1"},
        "附录C": {"附录C"},
    }
    for item in replacements:
        if item.get("status") not in {"pass", "warning"}:
            continue
        target = str(item.get("target") or item.get("caption") or "")
        if target.startswith(("图", "表", "附录")):
            prefixes.add(target)
        for key, values in target_to_prefixes.items():
            if key in target:
                prefixes.update(values)
    return prefixes


def _condition_rows(component: str, result: dict[str, Any], evaluation: list[dict[str, Any]]) -> list[list[Any]]:
    rows = result.get("beam_stress_results") or []
    by_case: dict[str, dict[str, float]] = {}
    for row in rows:
        if row.get("report_component_hint") != component and row.get("component") != component:
            continue
        case = str(row.get("load_case") or "").upper()
        stress_type = str(row.get("stress_type") or "").upper()
        value = row.get("value_mpa")
        if value is None:
            continue
        by_case.setdefault(case, {})
        current = by_case[case].get(stress_type)
        normalized = abs(float(value))
        if current is None or normalized > current:
            by_case[case][stress_type] = normalized

    eval_by_suffix = {
        str(item.get("check_id", "")).rsplit(".", 1)[-1]: item
        for item in evaluation
        if item.get("component") == component
    }

    def envelope(cases: set[str], stress_type: str) -> float:
        values = [data.get(stress_type, 0.0) for case, data in by_case.items() if case in cases]
        return max(values) if values else 0.0

    if not any(abs(value) > 1e-9 for data in by_case.values() for value in data.values()):
        return []

    def allowable(suffix: str, accident: bool = False) -> float | None:
        item = eval_by_suffix.get(suffix)
        if not item:
            return None
        base = item.get("allowable_value")
        if base is None:
            return None
        material_id = str(item.get("material_id") or "").lower()
        if accident and suffix not in {"support_tension_bending_combined", "support_compression_bending_combined"}:
            factor = 1.66 if material_id == "q235" else 1.5450422535211268
        else:
            factor = 1.0
        return float(base) * factor

    def stress_rows(label: str, cases: set[str], accident: bool) -> list[list[Any]]:
        values = {stress_type: envelope(cases, stress_type) for stress_type, _, _ in STRESS_LABELS}
        output: list[list[Any]] = []
        ratios: dict[str, float] = {}
        for stress_type, cn, suffix in STRESS_LABELS:
            allow = allowable(suffix, accident=accident)
            value = values[stress_type]
            ratio = value / allow if allow else None
            ratios[suffix] = ratio or 0.0
            output.append([label, cn, _fmt_num(value), _fmt_num(allow), _fmt_ratio(ratio)])
        tension_bending = ratios["support_tension"] + ratios["support_bending"]
        compression_bending = ratios["support_compression"] + ratios["support_bending"]
        output.append([label, "拉弯组合", _fmt_num(tension_bending, 2), "1", _fmt_ratio(tension_bending)])
        output.append([label, "压弯组合", _fmt_num(compression_bending, 2), "1", _fmt_ratio(compression_bending)])
        return output

    normal_cases = {"NORMAL", "UPSET", "B", "SL-1"}
    accident_cases = {"FAULTED", "D", "SL-2"}
    return stress_rows("正常异常", normal_cases, False) + stress_rows("事故", accident_cases, True)


def _delete_table_row(table, row_index: int) -> None:
    table._tbl.remove(table.rows[row_index]._tr)


def _fill_rows(table, rows: list[list[Any]], start_row: int = 1, *, trim_unused: bool = False) -> int:
    count = 0
    for offset, row_values in enumerate(rows):
        row_index = start_row + offset
        if row_index >= len(table.rows):
            break
        cells = table.rows[row_index].cells
        for column, value in enumerate(row_values[: len(cells)]):
            _set_cell_text(cells[column], value)
        count += 1
    if trim_unused:
        for row_index in range(len(table.rows) - 1, start_row + count - 1, -1):
            _delete_table_row(table, row_index)
        return count
    for row_index in range(start_row + count, len(table.rows)):
        for cell in table.rows[row_index].cells:
            _set_cell_text(cell, "")
    return count


def _mark_table_pending(table, *, numeric_start: int = 2) -> None:
    for row in table.rows[1:]:
        cells = row.cells
        for column in range(numeric_start, len(cells)):
            _set_cell_text(cells[column], "待确认")


def _weld_condition(item: dict[str, Any]) -> str:
    check_id = str(item.get("check_id") or "").lower()
    load_case = str(item.get("load_case") or item.get("case") or "").upper()
    if "fault" in check_id or "accident" in check_id or "sl-2" in check_id or load_case in {"FAULTED", "D", "SL-2"}:
        return "事故工况"
    return "异常工况"


def _fill_weld_evaluation_table(table, weld_eval_rows: list[dict[str, Any]]) -> int:
    if not weld_eval_rows:
        return 0
    header = [cell.text for cell in table.rows[0].cells]
    is_pair_row_layout = len(header) >= 7 and "剪应力" in header[1] and "等效" in header[4]
    if is_pair_row_layout:
        by_condition: dict[str, dict[str, dict[str, Any]]] = {}
        for item in weld_eval_rows:
            category = str(item.get("category") or "")
            if "剪" in category:
                kind = "shear"
            elif "等效" in category:
                kind = "equivalent"
            else:
                continue
            by_condition.setdefault(_weld_condition(item), {})[kind] = item
        rows: list[list[Any]] = []
        for condition in ("异常工况", "事故工况"):
            values = by_condition.get(condition) or {}
            shear = values.get("shear") or {}
            equivalent = values.get("equivalent") or {}
            if not shear and not equivalent:
                continue
            rows.append(
                [
                    condition,
                    _fmt_num(shear.get("calculation_value")),
                    _fmt_num(shear.get("allowable_value")),
                    _fmt_ratio(shear.get("ratio")),
                    _fmt_num(equivalent.get("calculation_value")),
                    _fmt_num(equivalent.get("allowable_value")),
                    _fmt_ratio(equivalent.get("ratio")),
                ]
            )
        return _fill_rows(table, rows, trim_unused=True)

    rows = [
        [
            _weld_condition(item),
            item.get("category"),
            _fmt_num(item.get("calculation_value")),
            _fmt_num(item.get("calculation_value")),
            _fmt_num(item.get("allowable_value")),
            _fmt_ratio(item.get("ratio")),
        ]
        for item in weld_eval_rows
        if "剪" in str(item.get("category") or "") or "等效" in str(item.get("category") or "")
    ]
    return _fill_rows(table, rows, trim_unused=True)


def _is_equivalent_weld_branch(scope: dict[str, Any] | None, evaluation: list[dict[str, Any]]) -> bool:
    requires = (scope or {}).get("requires") or {}
    if requires.get("cantilever_root_weld_equivalent_stress_table"):
        return True
    return any(str(item.get("check_id", "")).startswith("cantilever_root_weld_equivalent.") for item in evaluation)


def _ensure_equivalent_weld_table_layout(document: Document, table_index: int = 10) -> dict[str, Any]:
    if len(document.tables) <= table_index:
        return {"status": "warning", "message": "equivalent weld table target is missing"}
    table = document.tables[table_index]
    removed_columns = 0
    for grid_col in list(table._tbl.xpath("./w:tblGrid/w:gridCol"))[6:]:
        grid_col.getparent().remove(grid_col)
        removed_columns += 1
    for tr in table._tbl.tr_lst:
        for tc in list(tr.tc_lst)[6:]:
            tr.remove(tc)
        for tc in tr.tc_lst:
            tc_pr = tc.tcPr
            if tc_pr is None:
                continue
            for child in list(tc_pr):
                if child.tag in {qn("w:vMerge"), qn("w:gridSpan")}:
                    tc_pr.remove(child)
    headers = ["工况", "应力类型", "计算值(MPa)", "等效应力(MPa)", "许用值(MPa)", "应力比"]
    if len(table.rows[0].cells) < len(headers):
        return {"status": "warning", "message": "equivalent weld table has fewer than 6 columns after normalization"}
    for column, text in enumerate(headers):
        _set_cell_text(table.rows[0].cells[column], text)
    return {"status": "pass", "layout_normalized": True, "removed_columns": removed_columns}


def _equivalent_weld_rows(
    evaluation: list[dict[str, Any]],
    *,
    coefficient: float = WELD_EQUIVALENT_COEFFICIENT,
) -> list[list[Any]]:
    by_case_type: dict[tuple[str, str], dict[str, Any]] = {}
    stress_type_by_token = {
        "tension": "拉伸应力",
        "compression": "压缩应力",
        "bending": "弯曲应力",
        "shear": "剪切应力",
    }
    for item in evaluation:
        check_id = str(item.get("check_id") or "")
        if not check_id.startswith("cantilever_root_weld_equivalent."):
            continue
        if item.get("ratio") is None:
            continue
        case = "accident" if ".accident." in check_id else "normal_abnormal"
        token = check_id.rsplit(".", 1)[-1]
        if token not in stress_type_by_token:
            continue
        by_case_type[(case, token)] = item

    rows: list[list[Any]] = []
    for case, label in (("normal_abnormal", "正常/异常"), ("accident", "事故")):
        for token in ("tension", "compression", "bending", "shear"):
            item = by_case_type.get((case, token))
            if not item:
                continue
            calculation_value = item.get("calculation_value")
            try:
                equivalent_stress = float(calculation_value) / coefficient
            except (TypeError, ValueError, ZeroDivisionError):
                equivalent_stress = None
            rows.append(
                [
                    label,
                    stress_type_by_token[token],
                    _fmt_num(calculation_value, 3),
                    _fmt_num(equivalent_stress, 3),
                    _fmt_num(item.get("allowable_value"), 3),
                    _fmt_ratio_precise(item.get("ratio"), 3),
                ]
            )
    return rows


def _fill_equivalent_weld_table(
    document: Document,
    evaluation: list[dict[str, Any]],
    *,
    coefficient: float = WELD_EQUIVALENT_COEFFICIENT,
) -> dict[str, Any]:
    layout = _ensure_equivalent_weld_table_layout(document, 10)
    if layout.get("status") != "pass" or len(document.tables) <= 10:
        return {
            "target": "托臂根部焊缝评定结果（应力比）",
            "filled_rows": 0,
            "status": "warning",
            "message": layout.get("message", "equivalent weld table layout unavailable"),
        }
    rows = _equivalent_weld_rows(evaluation, coefficient=coefficient)
    table = document.tables[10]
    _ensure_table_rows(table, 1 + len(rows))
    filled = _fill_rows(table, rows, trim_unused=True)
    return {
        "target": "托臂根部焊缝评定结果（应力比）",
        "filled_rows": filled,
        "status": "pass" if filled else "warning",
        "equivalent_coefficient": coefficient,
        "layout_normalized": layout.get("layout_normalized", False),
        "removed_columns": layout.get("removed_columns", 0),
    }


def _fill_support_tables(
    document: Document,
    result: dict[str, Any],
    evaluation: list[dict[str, Any]],
    mode: str,
    scope: dict[str, Any] | None = None,
) -> list[dict[str, Any]]:
    replacements: list[dict[str, Any]] = []
    tables = document.tables
    if len(tables) <= 14:
        return [{"target": "result_tables", "status": "warning", "message": "template has fewer tables than expected"}]
    requires = (scope or {}).get("requires") or {}
    root_weld_required = bool(requires.get("cantilever_root_weld_eval"))
    equivalent_weld_branch = _is_equivalent_weld_branch(scope, evaluation)
    try:
        equivalent_weld_coefficient = float((scope or {}).get("cantilever_root_weld_equivalent_coefficient") or WELD_EQUIVALENT_COEFFICIENT)
    except (TypeError, ValueError):
        equivalent_weld_coefficient = WELD_EQUIVALENT_COEFFICIENT

    if mode == "steel_platform":
        mapping = [(8, "square_support", "支架方钢应力评定表"), (9, "cantilever_arm", "托臂应力评定表")]
    else:
        mapping = [(8, "mixed_beam_type_1", "支架应力评定表")]
    square_outer_width = (scope or {}).get("square_outer_width_mm")
    large_square_weld_principle = isinstance(square_outer_width, (int, float)) and float(square_outer_width) > 120.0
    if mode == "steel_platform":
        if large_square_weld_principle:
            mapping = [
                (8, "mixed_beam_type_1", "support stress evaluation"),
                (9, "square_support", "square support stress evaluation"),
            ]
        else:
            mapping = [
                (8, "mixed_beam_type_1", "support stress evaluation"),
                (9, "cantilever_arm", "cantilever/root weld stress evaluation"),
            ]

    for table_index, component, label in mapping:
        component_candidates = [component]
        if component == "mixed_beam_type_1":
            component_candidates.extend(["square_support", "support"])
        elif component == "cantilever_arm":
            component_candidates.append("mixed_beam_type_1")
        rows: list[list[Any]] = []
        chosen_component = component
        for candidate in component_candidates:
            rows = _condition_rows(candidate, result, evaluation)
            if rows:
                chosen_component = candidate
                break
        filled = _fill_rows(tables[table_index], rows)
        replacements.append({"target": label, "component": chosen_component, "filled_rows": filled, "status": "pass" if filled else "not_applicable"})

    if mode != "steel_platform":
        weld_force_rows = []
        for item in result.get("weld_force_results") or []:
            values = item.get("values") or item
            condition = "事故工况" if str(item.get("load_case", "")).upper() in {"FAULTED", "D", "SL-2"} else "异常工况"
            weld_force_rows.append(
                [
                    condition,
                    _fmt_num(_metric_value(values.get("fx")), 1),
                    _fmt_num(_metric_value(values.get("fy")), 1),
                    _fmt_num(_metric_value(values.get("fz")), 1),
                    _fmt_num(_metric_value(values.get("mx")), 1),
                    _fmt_num(_metric_value(values.get("my")), 1),
                    _fmt_num(_metric_value(values.get("mz")), 1),
                ]
            )
        if equivalent_weld_branch:
            replacements.append(
                {
                    "target": "托臂根部所受载荷",
                    "filled_rows": 0,
                    "status": "not_applicable",
                    "message": "square outer width <= 120 mm uses TMAXBEAMSTRESS equivalent weld-stress table; HF-FORCE root-load table is not required",
                }
            )
        elif weld_force_rows:
            filled = _fill_rows(tables[9], weld_force_rows)
            replacements.append({"target": "托臂根部所受载荷", "filled_rows": filled, "status": "pass" if filled else "warning"})
        elif not root_weld_required:
            replacements.append(
                {
                    "target": "托臂根部所受载荷",
                    "filled_rows": 0,
                    "status": "not_applicable",
                    "message": "analysis scope does not require cantilever root weld loads",
                }
            )
        else:
            _mark_table_pending(tables[9], numeric_start=1)
            replacements.append({"target": "托臂根部所受载荷", "filled_rows": 0, "status": "warning", "message": "no weld/root load rows in result.json"})

    weld_eval_rows = [item for item in evaluation if "weld" in str(item.get("check_id", "")).lower() and item.get("ratio") is not None]
    if equivalent_weld_branch:
        replacements.append(_fill_equivalent_weld_table(document, evaluation, coefficient=equivalent_weld_coefficient))
    elif weld_eval_rows:
        filled = _fill_weld_evaluation_table(tables[10], weld_eval_rows)
        replacements.append({"target": "焊缝评定结果", "filled_rows": filled, "status": "pass" if filled else "warning"})
    elif not root_weld_required:
        replacements.append(
            {
                "target": "焊缝评定结果",
                "filled_rows": 0,
                "status": "not_applicable",
                "message": "analysis scope does not require cantilever root weld evaluation",
            }
        )
    else:
        _mark_table_pending(tables[10], numeric_start=2)
        replacements.append({"target": "焊缝评定结果", "filled_rows": 0, "status": "warning", "message": "no confirmed weld stress rows in result.json"})

    foundation_rows = []
    for item in result.get("foundation_loads") or []:
        foundation_rows.append(
            [
                item.get("load_case"),
                _fmt_num(_metric_value(item.get("fx")), 1),
                _fmt_num(_metric_value(item.get("fy")), 1),
                _fmt_num(_metric_value(item.get("fz")), 1),
                _fmt_num(_metric_value(item.get("mx")), 1),
                _fmt_num(_metric_value(item.get("my")), 1),
                _fmt_num(_metric_value(item.get("mz")), 1),
            ]
        )
    filled = _fill_rows(tables[11], foundation_rows)
    replacements.append({"target": "支架基础载荷", "filled_rows": filled, "status": "pass" if filled else "warning"})

    bolt_rows = []
    for item in result.get("bolt_force_results") or []:
        condition = "事故工况" if str(item.get("load_case", "")).upper() in {"FAULTED", "D", "SL-2"} else "异常工况"
        values = item.get("values") or item
        bolt_rows.append(
            [
                "托盘与托臂",
                condition,
                _fmt_num(_metric_value(values.get("fx")), 1),
                _fmt_num(_metric_value(values.get("fy")), 1),
                _fmt_num(_metric_value(values.get("fz")), 1),
                _fmt_num(_metric_value(values.get("my")), 1),
                _fmt_num(_metric_value(values.get("mz")), 1),
            ]
        )
    filled = _fill_rows(tables[12], bolt_rows)
    replacements.append({"target": "支架连接螺栓载荷", "filled_rows": filled, "status": "pass" if filled else "warning"})

    bolt_eval_rows = [item for item in evaluation if "bolt" in str(item.get("check_id", "")).lower() and item.get("ratio") is not None]
    if bolt_eval_rows:
        rows = []
        by_case: dict[str, dict[str, Any]] = {}
        for item in bolt_eval_rows:
            key = "事故工况" if "fault" in str(item.get("check_id", "")).lower() or "_d_" in str(item.get("check_id", "")).lower() else "异常工况"
            by_case.setdefault(key, {})[item.get("category")] = item
        for condition in ("异常工况", "事故工况"):
            data = by_case.get(condition, {})
            shear = next((value for key, value in data.items() if "剪" in str(key)), {})
            tension = next((value for key, value in data.items() if "拉" in str(key)), {})
            combined = next((value for key, value in data.items() if "组合" in str(key)), {})
            rows.append(
                [
                    "托盘与托臂",
                    condition,
                    _fmt_num(shear.get("calculation_value")),
                    _fmt_num(shear.get("allowable_value")),
                    _fmt_num(tension.get("calculation_value")),
                    _fmt_num(tension.get("allowable_value")),
                    _fmt_ratio(combined.get("ratio")),
                ]
            )
        filled = _fill_rows(tables[13], rows)
        replacements.append({"target": "支架螺栓应力评定表", "filled_rows": filled, "status": "pass" if filled else "warning"})
    else:
        _mark_table_pending(tables[13], numeric_start=2)
        replacements.append({"target": "支架螺栓应力评定表", "filled_rows": 0, "status": "warning", "message": "no confirmed bolt stress rows in result.json"})

    modal_rows = []
    modes = result.get("modal_results") or []
    half = (len(modes) + 1) // 2
    for left, right in zip(modes[:half], modes[half:] + [{}] * half):
        modal_rows.append(
            [
                left.get("mode", ""),
                _fmt_num(left.get("frequency_hz"), 3),
                right.get("mode", ""),
                _fmt_num(right.get("frequency_hz"), 3) if right else "",
            ]
        )
    filled = _fill_rows(tables[14], modal_rows)
    replacements.append({"target": "模态频率表", "filled_rows": filled, "status": "pass" if filled else "warning"})
    return replacements


def _clear_paragraph(paragraph) -> None:
    p_pr = paragraph._p.pPr
    for child in list(paragraph._p):
        if child is not p_pr:
            paragraph._p.remove(child)


def _find_caption_index(document: Document, prefix: str) -> int | None:
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip().startswith(prefix):
            return index
    return None


def _find_caption_index_any(document: Document, prefixes: tuple[str, ...]) -> int | None:
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if any(text.startswith(prefix) for prefix in prefixes):
            return index
    return None


def _replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _replace_text_preserve_runs(paragraph, replacements: dict[str, str], regex_replacements: list[tuple[re.Pattern[str], str]] | None = None) -> bool:
    changed = False
    for run in paragraph.runs:
        text = run.text
        for old, new in replacements.items():
            if old and old in text:
                text = text.replace(old, new)
        for pattern, replacement in regex_replacements or []:
            text = pattern.sub(replacement, text)
        if text != run.text:
            run.text = text
            changed = True
    full_text = paragraph.text
    updated = full_text
    for old, new in replacements.items():
        if old:
            updated = updated.replace(old, new)
    for pattern, replacement in regex_replacements or []:
        updated = pattern.sub(replacement, updated)
    if updated != full_text:
        _replace_paragraph_text(paragraph, updated)
        changed = True
    return changed


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _iter_document_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        yield from _iter_table_paragraphs(table)
    for section in document.sections:
        for part in (section.header, section.first_page_header, section.even_page_header, section.footer, section.first_page_footer, section.even_page_footer):
            for paragraph in part.paragraphs:
                yield paragraph
            for table in part.tables:
                yield from _iter_table_paragraphs(table)


def _replace_report_identity(document: Document, input_payload: dict[str, Any], report_id: str) -> dict[str, Any]:
    project = input_payload.get("project") or {}
    building = str(project.get("building") or project.get("area") or "").strip()
    building_label = f"{building}厂房" if building and "厂房" not in building else building
    replacements: dict[str, str] = {}
    if building_label:
        replacements.update(
            {
                "NR厂房": building_label,
                "NR 厂房": f"{building_label} ",
                "NR Building": f"{building} Building",
            }
        )
    regex_replacements = [(re.compile(r"\b\d{5}NI-LXSJ\d+\b"), report_id)]
    changed = 0
    for paragraph in _iter_document_paragraphs(document):
        if _replace_text_preserve_runs(paragraph, replacements, regex_replacements):
            changed += 1
    return {
        "target": "报告号/厂房页眉页脚",
        "status": "pass",
        "changed_paragraphs": changed,
        "report_id": report_id,
        "building": building,
    }


def _image_path(job_dir: Path, *names: str) -> Path | None:
    for name in names:
        candidates = [job_dir / name]
        if not Path(name).parts or Path(name).parent == Path("."):
            stem = Path(name).stem
            candidates.extend(
                [
                    job_dir / "figures" / f"{stem}.png",
                    job_dir / "figures" / f"{stem}.PNG",
                    job_dir / "figures" / f"{stem}.bmp",
                    job_dir / "figures" / f"{stem}.BMP",
                ]
            )
        for path in candidates:
            if path.exists():
                return path
    return None


def _figure_path_from_manifest(
    job_dir: Path,
    figures: list[dict[str, Any]],
    *,
    component_scope: str | None = None,
    figure_type: str | None = None,
    names: tuple[str, ...] = (),
) -> Path | None:
    normalized_names = {name.upper() for name in names}
    for figure in figures:
        if component_scope and figure.get("component_scope") != component_scope:
            continue
        if figure_type and figure.get("figure_type") != figure_type:
            continue
        source_name = Path(str(figure.get("source_file") or "")).name.upper()
        target_name = Path(str(figure.get("target_file") or figure.get("path") or "")).name.upper()
        if normalized_names and source_name not in normalized_names and target_name not in normalized_names:
            continue
        target = job_dir / str(figure.get("target_file") or figure.get("path") or "")
        if target.exists():
            return target
        source = job_dir / str(figure.get("source_file") or "")
        if source.exists():
            return source
    return None


def _paragraph_has_image(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing") or paragraph._p.xpath(".//w:pict"))


def _target_image_paragraph_index(document: Document, caption_index: int) -> int:
    for index in range(caption_index - 1, max(-1, caption_index - 7), -1):
        paragraph = document.paragraphs[index]
        if _paragraph_has_image(paragraph) or not paragraph.text.strip():
            return index
    return max(0, caption_index - 1)


def _insert_image_before_caption(
    document: Document,
    caption_prefix: str,
    image_path: Path,
    width_inches: float = 5.8,
    *,
    alternate_prefixes: tuple[str, ...] = (),
    ensure_caption_text: str | None = None,
) -> dict[str, Any]:
    prefixes = (caption_prefix, *alternate_prefixes)
    index = _find_caption_index_any(document, prefixes)
    caption_added = False
    if index is None:
        if not ensure_caption_text:
            return {"caption": caption_prefix, "image": str(image_path), "status": "warning", "message": "caption not found"}
        document.add_paragraph("")
        caption = document.add_paragraph(ensure_caption_text)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        index = len(document.paragraphs) - 1
        caption_added = True
    target_index = _target_image_paragraph_index(document, index)
    paragraph = document.paragraphs[target_index]
    _clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))
    caption_paragraph = document.paragraphs[index]
    caption_paragraph.paragraph_format.keep_together = True
    _mark_paragraph_red(caption_paragraph)
    return {
        "caption": caption_prefix,
        "matched_caption": caption_paragraph.text.strip(),
        "image": image_path.name,
        "status": "pass",
        "caption_added_to_template_copy": caption_added,
        "title_marked_red": True,
    }


def _remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _paragraph_has_page_break(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:br[@w:type='page']") or paragraph._p.xpath("./w:pPr/w:pageBreakBefore"))


def _remove_empty_page_breaks_inside_appendices(document: Document) -> dict[str, Any]:
    appendix_start = None
    for index, paragraph in enumerate(document.paragraphs):
        style_name = (paragraph.style.name or "").lower()
        if style_name.startswith("toc"):
            continue
        if paragraph.text.strip().startswith(("附录A", "附录B", "附录C")):
            appendix_start = index
            break
    if appendix_start is None:
        return {"target": "附录分页符", "status": "not_applicable", "removed": 0}
    removed: list[dict[str, Any]] = []
    for index, paragraph in list(enumerate(document.paragraphs))[appendix_start + 1 :]:
        if paragraph.text.strip() or _paragraph_has_image(paragraph) or not _paragraph_has_page_break(paragraph):
            continue
        next_text = ""
        for next_paragraph in document.paragraphs[index + 1 :]:
            next_text = next_paragraph.text.strip()
            if next_text or _paragraph_has_image(next_paragraph):
                break
        removed.append({"paragraph_index": index, "next_text": next_text[:40]})
        _remove_paragraph(paragraph)
    return {
        "target": "附录分页符",
        "status": "pass",
        "removed": len(removed),
        "removed_paragraphs": removed,
        "message": "Removed empty page-break-only paragraphs inside appendices from the generated copy to avoid blank appendix pages.",
    }


def _remove_table(table) -> None:
    element = table._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _remove_body_paragraph_range(document: Document, start_text: str, end_text: str) -> int:
    start_index = None
    end_index = None
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if start_index is None and text == start_text:
            start_index = index
            continue
        if start_index is not None and text == end_text:
            end_index = index
            break
    if start_index is None:
        return 0
    if end_index is None:
        end_index = len(document.paragraphs)
    count = 0
    for paragraph in list(document.paragraphs[start_index:end_index]):
        _remove_paragraph(paragraph)
        count += 1
    return count


def _remove_not_applicable_weld_section(
    document: Document,
    scope: dict[str, Any] | None,
    *,
    has_weld_evaluation_rows: bool = False,
) -> dict[str, Any]:
    requires = (scope or {}).get("requires") or {}
    if requires.get("cantilever_root_weld_eval") or has_weld_evaluation_rows:
        return {"target": "6.2 焊缝的评定", "status": "not_applicable", "message": "weld section is required or has confirmed weld rows"}
    removed_tables = 0
    # Tables 9 and 10 in both templates are the root-weld load/evaluation tables.
    for table_index in (10, 9):
        if len(document.tables) > table_index:
            _remove_table(document.tables[table_index])
            removed_tables += 1
    removed_paragraphs = _remove_body_paragraph_range(document, "焊缝的评定", "基础载荷")
    return {
        "target": "6.2 焊缝的评定",
        "status": "pass",
        "removed_tables": removed_tables,
        "removed_paragraphs": removed_paragraphs,
        "message": "cantilever root weld section removed because current scope uses cantilever cloud figures instead of weld-principle evaluation",
    }


def _remove_matching_paragraphs(document: Document, predicates: list) -> int:
    removed = 0
    for paragraph in list(document.paragraphs):
        text = paragraph.text.strip()
        if any(predicate(text) for predicate in predicates):
            _remove_paragraph(paragraph)
            removed += 1
    return removed


def _replace_first_paragraph_starting_with(document: Document, prefix: str, text: str) -> bool:
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            _replace_paragraph_text(paragraph, text)
            return True
    return False


def _adapt_equivalent_weld_section(
    document: Document,
    scope: dict[str, Any] | None,
    evaluation: list[dict[str, Any]],
) -> dict[str, Any]:
    if not _is_equivalent_weld_branch(scope, evaluation):
        return {"target": "6.2 等效焊缝评定", "status": "not_applicable"}
    removed_tables = 0
    if len(document.tables) > 9:
        header = " ".join(cell.text.strip() for cell in document.tables[9].rows[0].cells)
        if "FX" in header and "FZ" in header and "MZ" in header:
            _remove_table(document.tables[9])
            removed_tables = 1
    removed_paragraphs = _remove_matching_paragraphs(
        document,
        [
            lambda text: text.startswith("托臂与托臂底板之间焊缝的焊角高度"),
            lambda text: text.startswith("图6.1 托臂根部焊缝截面投影形状"),
            lambda text: text.startswith("托臂根部所受载荷如下表所示"),
            lambda text: text.startswith("表6-2 托臂根部所受载荷"),
            lambda text: text.startswith("根据上述载荷计算得到各个工况下的焊缝剪应力"),
        ],
    )
    intro_changed = _replace_first_paragraph_starting_with(
        document,
        "支架的托臂与托臂底板之间",
        "支架的托臂与托臂底板之间采用角焊进行连接。方钢外边长不大于120mm时，按托臂根部梁单元应力结果并采用等效系数0.526进行焊缝等效评定，评定结果如下表所示。",
    )
    title_changed = _replace_first_paragraph_starting_with(
        document,
        "表6-3 托臂根部焊缝评定结果",
        "表6-2 托臂根部焊缝评定结果（应力比）",
    )
    marked_titles = []
    for prefix in ("焊缝的评定", "表6-2 托臂根部焊缝评定结果"):
        marked = _mark_first_title_red(document, prefix)
        if marked:
            marked_titles.append(marked)
    return {
        "target": "6.2 等效焊缝评定",
        "status": "pass",
        "removed_tables": removed_tables,
        "removed_paragraphs": removed_paragraphs,
        "intro_changed": intro_changed,
        "table_title_changed": title_changed,
        "marked_titles": marked_titles,
        "message": "root-load table removed for <=120 mm equivalent weld-stress branch; table values come from evaluation_summary.json",
    }


def _appendix_c_heading_index(document: Document) -> int | None:
    for index, paragraph in enumerate(document.paragraphs):
        style_name = (paragraph.style.name or "").lower()
        if style_name.startswith("toc"):
            continue
        if paragraph.text.strip().startswith("附录C"):
            return index
    return None


def _replace_appendix_c_tail(document: Document, lines: list[str]) -> None:
    heading_index = _appendix_c_heading_index(document)
    if heading_index is None:
        heading = document.add_paragraph()
        heading.style = "Heading 1"
        heading_index = len(document.paragraphs) - 1
    for paragraph in list(document.paragraphs[heading_index + 1 :]):
        _remove_paragraph(paragraph)
    _replace_paragraph_text(document.paragraphs[heading_index], lines[0])
    for line in lines[1:]:
        document.add_paragraph(line)


def _adapt_appendix_c(document: Document, scope: dict[str, Any] | None) -> dict[str, Any]:
    mode = (scope or {}).get("appendix_c_mode")
    heading_index = _appendix_c_heading_index(document)
    current_heading = document.paragraphs[heading_index].text.strip() if heading_index is not None else ""
    if mode == "cantilever_stress_cloud":
        if current_heading.startswith("附录C：托臂应力图"):
            return {"target": "附录C", "mode": mode, "status": "pass", "message": "template appendix C already matches cantilever stress figures"}
        lines = [
            "附录C：托臂应力图",
            "",
            "图C-1 异常工况下托臂轴向应力分布",
            "",
            "图C-2 异常工况下托臂轴向应力分布",
            "",
            "图C-3 异常工况下托臂弯曲应力分布",
            "",
            "图C-4 异常工况下托臂剪切应力分布",
            "",
            "图C-5 事故工况下托臂轴向应力分布",
            "",
            "图C-6 事故工况下托臂轴向应力分布",
            "",
            "图C-7 事故工况下托臂弯曲应力分布",
            "",
            "图C-8 事故工况下托臂剪切应力分布",
        ]
        _replace_appendix_c_tail(document, lines)
        return {"target": "附录C", "mode": mode, "status": "pass", "message": "appendix C configured for cantilever stress figures"}
    if mode == "weld_evaluation_principle":
        if current_heading.startswith("附录C：焊缝评定原理"):
            return {"target": "附录C", "mode": mode, "status": "pass", "message": "template appendix C already matches weld evaluation principle"}
        lines = [
            "附录C：焊缝评定原理",
            "焊缝的计算原理如下，焊缝最外缘为焊缝截面上最危险位置, 此处以一个矩形底板四周焊缝（示意图见图C-1）为例：",
            "",
            "图C-1 环板焊缝示意图",
            "（灰色区域为矩形底板，蓝色区域为焊缝有效焊喉截面）",
            "",
            "由上述值可以得到：",
            "",
            "然后评定上述剪应力，等效应力。",
            "注：评定中用的载荷FX/FY/FZ/MX/MY/MZ是环板受到相对于中心点O的合力。按照《钢结构设计规范GB50017-2017》中规定的按照能量强度理论折算应力，设计时使其不超过按角焊缝受力时的焊缝熔敷金属抗拉强度设计值。",
        ]
        _replace_appendix_c_tail(document, lines)
        return {"target": "附录C", "mode": mode, "status": "pass", "message": "appendix C configured for weld evaluation principle"}
    return {"target": "附录C", "mode": mode or "unknown", "status": "not_applicable"}


def _replace_figures(
    document: Document,
    job_dir: Path,
    figures: list[dict[str, Any]],
    mode: str,
    scope: dict[str, Any] | None = None,
) -> list[dict[str, Any]]:
    replacements: list[dict[str, Any]] = []
    model = _figure_path_from_manifest(
        job_dir,
        figures,
        component_scope="whole_model",
        figure_type="model",
        names=("SHITI.PNG", "SHITI.BMP"),
    ) or _image_path(job_dir, "SHITI.PNG", "SHITI.bmp")
    arm = _figure_path_from_manifest(
        job_dir,
        figures,
        component_scope="cantilever_model",
        figure_type="model",
        names=("TBMODEL.PNG", "TBMODEL.BMP", "TUOBI_MODEL.PNG", "TUOBI_MODEL.BMP", "TUOBI_MO.PNG"),
    ) or _image_path(job_dir, "TBMODEL.PNG", "TBMODEL.bmp", "TUOBI_MODEL.PNG", "TUOBI_MODEL.bmp", "TUOBI_MO.PNG")
    if model:
        replacements.append(_insert_image_before_caption(document, "图5.1", model, 5.8, ensure_caption_text="图5.1 S2 支架有限元模型"))
    else:
        replacements.append({"caption": "图5.1", "status": "warning", "message": "whole-model figure SHITI is missing"})
    if arm:
        replacements.append(_insert_image_before_caption(document, "图5.2", arm, 5.8, ensure_caption_text="图5.2 托臂有限元模型"))
    else:
        replacements.append(
            {
                "caption": "图5.2",
                "status": "warning",
                "message": "dedicated cantilever model figure TBMODEL is missing; SHITI is not reused for Fig. 5.2",
            }
        )

    for mode_index in range(1, 5):
        path = _image_path(job_dir, f"MOTAI-{mode_index}.PNG", f"MOTAI-{mode_index}.bmp")
        if path:
            replacements.append(
                _insert_image_before_caption(
                    document,
                    f"图A-{mode_index}",
                    path,
                    5.8,
                    ensure_caption_text=f"图A-{mode_index} 第{mode_index}阶模态振型图",
                )
            )

    case_order = {"B": 0, "UPSET": 0, "SL-1": 0, "D": 1, "FAULTED": 1, "SL-2": 1}
    stress_order = {"SDIR1": 0, "SDIR_TENSION": 0, "SDIR2": 1, "SDIR_COMPRESSION": 1, "SBEND": 2, "SHEAR": 3}

    def figure_order(item: dict[str, Any]) -> tuple[int, int, str]:
        return (
            case_order.get(str(item.get("case") or item.get("load_case") or "").upper(), 99),
            stress_order.get(str(item.get("stress_type") or "").upper(), 99),
            str(item.get("target_file") or item.get("path") or ""),
        )

    appendix_b = [figure for figure in figures if figure.get("appendix") == "B" and figure.get("component_scope") == "square_support"]
    appendix_b = sorted(appendix_b, key=figure_order)
    for index, figure in enumerate(appendix_b, start=1):
        path = job_dir / (figure.get("target_file") or figure.get("path") or "")
        if path.exists():
            replacements.append(
                _insert_image_before_caption(
                    document,
                    f"图B-{index}",
                    path,
                    5.8,
                    ensure_caption_text=f"图B-{index} 方钢应力图",
                )
            )

    requires = (scope or {}).get("requires") or {}
    if bool(requires.get("appendix_c_cantilever_figures")):
        appendix_c = [figure for figure in figures if figure.get("appendix") == "C" and figure.get("component_scope") == "cantilever_arm"]
        appendix_c = sorted(appendix_c, key=figure_order)
        for index, figure in enumerate(appendix_c, start=1):
            path = job_dir / (figure.get("target_file") or figure.get("path") or "")
            if path.exists():
                replacements.append(
                    _insert_image_before_caption(
                        document,
                        f"图C-{index}",
                        path,
                        5.8,
                        ensure_caption_text=f"图C-{index} 托臂应力图",
                    )
                )
    return replacements


def _points_from_spectrum_file(job_dir: Path, input_payload: dict[str, Any]) -> dict[str, list[dict[str, float]]] | None:
    points_path = job_dir / "spectrum_points.json"
    points = _read_json(points_path, {})
    load_steps = points.get("load_steps") if isinstance(points, dict) else None
    if not load_steps:
        metadata = input_payload.get("metadata") or {}
        source = metadata.get("static_acceleration_source") or {}
        workbook = source.get("workbook") or (input_payload.get("spectrum") or {}).get("spectrum_file")
        if not workbook:
            return None
        try:
            from core.spectra.response_spectrum_writer import _curve_payload
            from core.spectra.static_coefficients import DAMPING_BY_LEVEL, _curve_at_elevation, _read_segmented_sheet, resolve_segmented_spectrum_sheet
            from core.spectra.workbook_envelope import _envelope_curves

            workbook_path = Path(workbook)
            project = input_payload.get("project") or {}
            metadata = input_payload.get("metadata") or {}
            building = project.get("building") or project.get("area") or ""
            project_code = str(project.get("project_code") or "")
            elevation = float(project.get("elevation") or 0.0)
            static_elevations = _static_report_spectrum_elevations(input_payload) or [
                float(item)
                for item in (metadata.get("static_elevation_candidates") or [])
                if isinstance(item, int | float) or re.match(r"^[-+]?\d+(?:\.\d+)?$", str(item).strip())
            ]
            elevations = sorted(set(static_elevations or [elevation]))
            sheet = resolve_segmented_spectrum_sheet(workbook_path, building=building, project_code=project_code)
            curves = _read_segmented_sheet(workbook_path, sheet)
            load_steps = []
            for level in ("SL-1", "SL-2"):
                damping = DAMPING_BY_LEVEL[level]
                x_curves = [_curve_at_elevation(curves, level=level, direction="X", elevation=item, damping=damping) for item in elevations]
                y_curves = [_curve_at_elevation(curves, level=level, direction="Y", elevation=item, damping=damping) for item in elevations]
                z_curves = [_curve_at_elevation(curves, level=level, direction="Z", elevation=item, damping=damping) for item in elevations]
                horizontal = _envelope_curves([*x_curves, *y_curves], "+".join(curve.source_ref for curve in [*x_curves, *y_curves]) + ":horizontal_elevation_envelope")
                vertical = _envelope_curves(z_curves, "+".join(curve.source_ref for curve in z_curves) + ":vertical_elevation_envelope")
                load_steps.append({"level": level, "direction": "Z", "points": _curve_payload(vertical)})
                load_steps.append({"level": level, "direction": "X", "points": _curve_payload(horizontal)})
        except Exception:
            return None
    curves_by_key: dict[str, list[dict[str, float]]] = {}
    horizontal_candidates: dict[str, list[list[dict[str, float]]]] = {"SL-1": [], "SL-2": []}

    def merge_points(point_sets: list[list[dict[str, float]]]) -> list[dict[str, float]]:
        merged: dict[float, dict[str, float]] = {}
        for point_set in point_sets:
            for point in point_set:
                try:
                    frequency = round(float(point.get("frequency_hz")), 10)
                    acceleration = float(point.get("acceleration_g"))
                except (TypeError, ValueError):
                    continue
                current = merged.get(frequency)
                if current is None or abs(acceleration) > abs(float(current["acceleration_g"])):
                    merged[frequency] = {
                        "frequency_hz": float(point.get("frequency_hz")),
                        "acceleration_g": acceleration,
                    }
        return [merged[key] for key in sorted(merged)]

    for item in load_steps:
        level = item.get("level")
        direction = item.get("direction")
        if level not in {"SL-1", "SL-2"}:
            continue
        points = item.get("points") or []
        if direction == "Z":
            curves_by_key[f"{level}_vertical"] = merge_points([curves_by_key.get(f"{level}_vertical", []), points])
        elif direction in {"X", "Y"}:
            horizontal_candidates[level].append(points)
        else:
            horizontal_candidates[level].append(points)
    for level, point_sets in horizontal_candidates.items():
        if point_sets:
            curves_by_key[f"{level}_horizontal"] = merge_points(point_sets)
    return curves_by_key


def _ensure_table_rows(table, required_rows: int) -> None:
    while len(table.rows) < required_rows:
        table._tbl.append(deepcopy(table.rows[-1]._tr))


def _clear_table_tail(table, start_row: int) -> None:
    for row in table.rows[start_row:]:
        for cell in row.cells:
            _set_cell_text(cell, "")


def _fill_spectrum_table(document: Document, input_payload: dict[str, Any], job_dir: Path) -> dict[str, Any]:
    if len(document.tables) <= 2:
        return {"target": "表3-1", "status": "warning", "message": "spectrum table not found"}
    project = input_payload.get("project") or {}
    metadata = input_payload.get("metadata") or {}
    building = project.get("building") or project.get("area") or "厂房"
    elevation = project.get("elevation")
    caption_index = _find_caption_index(document, "表3-1")
    if caption_index is not None:
        static_elevations = _static_report_spectrum_elevations(input_payload)
        if not static_elevations:
            for item in metadata.get("static_elevation_candidates") or []:
                try:
                    static_elevations.append(float(item))
                except (TypeError, ValueError):
                    continue
        if static_elevations:
            elevations = sorted(set(static_elevations))
            elevation_text = "、".join(f"{_fmt_num(item, 2).rstrip('0').rstrip('.')}m" for item in elevations)
            suffix = "包络楼层反应谱" if len(elevations) > 1 else "楼层反应谱"
            _replace_paragraph_text(document.paragraphs[caption_index], f"表3-1 {building}{elevation_text}{suffix}")
        else:
            elevation_text = _fmt_num(elevation, 2).rstrip("0").rstrip(".") if elevation is not None else ""
            _replace_paragraph_text(document.paragraphs[caption_index], f"表3-1 {building}{elevation_text}m楼层反应谱")
        _mark_paragraph_red(document.paragraphs[caption_index])

    curves = _points_from_spectrum_file(job_dir, input_payload)
    if not curves:
        return {"target": "表3-1", "status": "warning", "message": "spectrum points unavailable; caption updated only"}
    table = document.tables[2]
    keys = ["SL-1_vertical", "SL-1_horizontal", "SL-2_vertical", "SL-2_horizontal"]
    curve_lengths = {key: len(curves.get(key, [])) for key in keys}
    max_rows = max(curve_lengths.values()) if curve_lengths else 0
    _ensure_table_rows(table, 3 + max_rows)
    for index in range(max_rows):
        cells = table.rows[index + 3].cells
        values = []
        for key in keys:
            curve = curves.get(key, [])
            if index < len(curve):
                point = curve[index]
                values.extend([_fmt_num(point.get("frequency_hz"), 3), _fmt_num(point.get("acceleration_g"), 4)])
            else:
                values.extend(["", ""])
        for column, value in enumerate(values[: len(cells)]):
            _set_cell_text(cells[column], value)
    _clear_table_tail(table, 3 + max_rows)
    return {
        "target": "表3-1",
        "status": "pass" if max_rows else "warning",
        "filled_rows": max_rows,
        "curve_lengths": curve_lengths,
        "table_rows_after": len(table.rows),
    }


def build_report_from_template(job_dir: Path | str, *, template_dir: Path | str = TEMPLATE_DIR, output_path: Path | str | None = None) -> dict[str, Any]:
    job_dir = Path(job_dir)
    result = _read_json(job_dir / "result.json", {})
    input_payload = _read_json(job_dir / "input.json", {})
    figures = _read_json(job_dir / "figures_manifest.json", [])
    evaluation = result.get("evaluation_summary") or _read_json(job_dir / "evaluation_summary.json", [])
    if isinstance(figures, dict):
        figures = figures.get("figures") or figures.get("items") or []
    if not isinstance(figures, list):
        figures = []
    if not isinstance(evaluation, list):
        evaluation = []

    selection = choose_report_template(input_payload, template_dir)
    template_path = Path(selection["template"])
    if not template_path.exists():
        raise FileNotFoundError(f"report template not found: {template_path}")

    document = Document(str(template_path))
    scope = classify_scope_from_input(input_payload)
    replacements: list[dict[str, Any]] = []
    report_id = _report_id(job_dir, input_payload)
    replacements.append(_replace_report_identity(document, input_payload, report_id))
    replacements.append(_fill_spectrum_table(document, input_payload, job_dir))
    replacements.extend(_fill_support_tables(document, result, evaluation, selection["mode"], scope))
    replacements.append(_adapt_equivalent_weld_section(document, scope, evaluation))
    replacements.append(_adapt_appendix_c(document, scope))
    replacements.extend(_replace_figures(document, job_dir, figures, selection["mode"], scope))
    replacements.append(
        _remove_not_applicable_weld_section(
            document,
            scope,
            has_weld_evaluation_rows=any(
                "weld" in str(item.get("check_id", "")).lower() and item.get("ratio") is not None
                for item in evaluation
            ),
        )
    )
    replacements.append(_remove_empty_page_breaks_inside_appendices(document))
    marked_titles = _mark_titles_red(document, _replacement_title_prefixes(replacements))
    if marked_titles:
        replacements.append(
            {
                "target": "修改/注入内容标题标红",
                "status": "pass",
                "marked_titles": sorted(set(marked_titles)),
                "message": "Only section titles, table titles and figure captions touched by template injection are marked red; body text and table values are not colored.",
            }
        )

    output = Path(output_path) if output_path else job_dir / f"{report_id}.docx"
    document.save(output)

    compatible_output = job_dir / "report.docx"
    if output.resolve() != compatible_output.resolve():
        shutil.copyfile(output, compatible_output)

    audit = {
        "status": "pass" if all(item.get("status") in {"pass", "not_applicable"} for item in replacements) else "warning",
        "job_id": job_dir.name,
        "template_mode": selection["mode"],
        "template": str(template_path),
        "report_file": output.name,
        "compatible_report_file": compatible_output.name,
        "replacements": replacements,
        "policy": "Template content and structure are preserved; table values come from result.json/input.json and figures come from ANSYS output files recorded in figures_manifest.json.",
    }
    (job_dir / "template_report_audit.json").write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    return audit
