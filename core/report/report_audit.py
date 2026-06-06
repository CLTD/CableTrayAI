from __future__ import annotations

import hashlib
import json
from pathlib import Path

from docx import Document

from core.report.template_mapper import validate_report_mapping


REQUIRED_HEADINGS = [
    "1 概述",
    "2 结构说明",
    "2.1 材料参数",
    "2.2 结构设计参数",
    "3 工况",
    "3.1 载荷工况",
    "3.2 工况组合",
    "4 规范要求",
    "5 计算方法、程序和计算模型",
    "6 结果评定",
    "6.1 支架的评定",
    "6.2 焊缝的评定",
    "6.3 基础载荷",
    "6.4 螺栓的评定",
    "7 结论",
    "参考资料",
    "附录A 模态分析结果",
    "附录B 支架应力图",
    "附录C 焊缝评定原理",
]


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _max_ratio(evaluation: list[dict]) -> float | None:
    ratios = [float(item["ratio"]) for item in evaluation if item.get("ratio") is not None]
    return max(ratios) if ratios else None


def _expected_conclusion(evaluation: list[dict]) -> str:
    if any(item.get("pass_fail") == "不满足" for item in evaluation):
        return "不满足"
    if any(item.get("formula_status") == "unconfirmed_todo" or item.get("pass_fail") == "待确认" for item in evaluation):
        return "待确认"
    return "满足"


def audit_report(job_dir: Path | str) -> dict:
    job_dir = Path(job_dir)
    result_path = job_dir / "result.json"
    figures_path = job_dir / "figures_manifest.json"
    input_path = job_dir / "input.json"
    report_path = job_dir / "report.docx"
    result = json.loads(result_path.read_text(encoding="utf-8"))
    input_payload = json.loads(input_path.read_text(encoding="utf-8")) if input_path.exists() else {}
    evaluation = result.get("evaluation_summary", [])
    figures = json.loads(figures_path.read_text(encoding="utf-8"))

    mapping_checks = validate_report_mapping(result, input_payload, figures)
    figure_checks = []
    for figure in figures:
        target = job_dir / figure.get("target_file", figure.get("path", ""))
        figure_checks.append(
            {
                "figure_id": figure.get("figure_id"),
                "target_file": figure.get("target_file", figure.get("path")),
                "exists": target.exists(),
            }
        )
    conclusion = _expected_conclusion(evaluation)
    max_ratio = _max_ratio(evaluation)
    status = "pass"
    if not report_path.exists() or report_path.stat().st_size == 0:
        status = "fail"
    if any(check["status"] != "pass" for check in mapping_checks):
        status = "fail"
    if any(not check["exists"] for check in figure_checks):
        status = "fail"

    headings: list[str] = []
    table_count = 0
    if report_path.exists():
        document = Document(str(report_path))
        headings = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.style.name.startswith("Heading") and paragraph.text.strip()]
        table_count = len(document.tables)
    missing_headings = [heading for heading in REQUIRED_HEADINGS if heading not in headings]
    if missing_headings:
        status = "fail"
    if table_count < 7:
        status = "fail"

    audit = {
        "status": status,
        "report_file": report_path.name,
        "result_sha256": _sha256(result_path),
        "report_sha256": _sha256(report_path) if report_path.exists() else None,
        "mapping_checks": mapping_checks,
        "figure_checks": figure_checks,
        "structure_checks": {
            "table_count": table_count,
            "minimum_table_count": 7,
            "table_count_status": "pass" if table_count >= 7 else "fail",
            "missing_headings": missing_headings,
            "heading_status": "pass" if not missing_headings else "fail",
        },
        "result_tables": {
            "modal_rows": len(result.get("modal_results", [])),
            "support_eval_rows": len([item for item in evaluation if item["check_id"].startswith("support_")]),
            "weld_eval_rows": len([item for item in evaluation if item["check_id"].startswith("weld") or "_weld_" in item["check_id"]]),
            "bolt_eval_rows": len([item for item in evaluation if "_bolt_" in item["check_id"]]),
            "foundation_load_rows": len(result.get("foundation_loads", [])),
            "figure_rows": len(figures),
        },
        "conclusion_check": {
            "max_ratio": max_ratio,
            "expected_conclusion": conclusion,
            "unconfirmed_formula_items": [
                item["check_id"] for item in evaluation if item.get("formula_status") == "unconfirmed_todo" or item.get("pass_fail") == "待确认"
            ],
            "failed_items_reported": [item["check_id"] for item in evaluation if item.get("pass_fail") == "不满足"],
        },
        "source_refs": sorted({str(item.get("source_ref")) for item in evaluation if item.get("source_ref")}),
        "data_source_policy": "All report numeric values are written from result.json/input.json; figures are referenced from figures_manifest.json.",
    }
    (job_dir / "report_audit.json").write_text(
        json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return audit
