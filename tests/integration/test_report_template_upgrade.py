from pathlib import Path

from docx import Document


def test_fixed_report_templates_are_readable_docx_files():
    template_paths = [
        Path("templates/report/non_steel_platform_report_template.docx"),
        Path("templates/report/steel_platform_report_template.docx"),
    ]

    for template_path in template_paths:
        assert template_path.exists()
        document = Document(str(template_path))
        assert len(document.paragraphs) > 0
        assert len(document.tables) > 0


def test_report_templates_are_not_empty_placeholders():
    for template_path in Path("templates/report").glob("*platform_report_template.docx"):
        assert template_path.stat().st_size > 100_000
