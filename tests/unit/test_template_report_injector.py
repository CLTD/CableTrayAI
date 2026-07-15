from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK

from types import SimpleNamespace

from core.report.template_injector import choose_report_template
from core.report.template_injector import (
    _equivalent_weld_rows,
    _paragraph_has_page_break,
    _points_from_spectrum_file,
    _remove_empty_page_breaks_inside_appendices,
    _static_report_spectrum_elevations,
)


def test_choose_report_template_returns_existing_static_template():
    selection = choose_report_template({"metadata": {"analysis_method": "static"}})

    assert selection["mode"] == "steel_platform"
    assert Path(selection["template"]).exists()


def test_choose_report_template_returns_existing_response_spectrum_template():
    selection = choose_report_template({"metadata": {"analysis_method": "response_spectrum"}})

    assert selection["mode"] == "non_steel_platform"
    assert Path(selection["template"]).exists()


def test_equivalent_weld_rows_use_tmax_values_and_equivalent_coefficient():
    evaluation = [
        {
            "check_id": "cantilever_root_weld_equivalent.normal_abnormal.tension",
            "category": "拉伸应力",
            "calculation_value": 0.9791529,
            "allowable_value": 159.75,
            "ratio": 0.011652628572448634,
        },
        {
            "check_id": "cantilever_root_weld_equivalent.normal_abnormal.bending",
            "category": "弯曲应力",
            "calculation_value": 65.650667,
            "allowable_value": 234.3,
            "ratio": 0.5326980537447521,
        },
        {
            "check_id": "cantilever_root_weld_equivalent.accident.shear",
            "category": "剪切应力",
            "calculation_value": 14.9224935,
            "allowable_value": 219.396,
            "ratio": 0.12930846280562736,
        },
    ]

    rows = _equivalent_weld_rows(evaluation, coefficient=0.526)

    assert rows == [
        ["正常/异常", "拉伸应力", "0.979", "1.862", "159.750", "0.012"],
        ["正常/异常", "弯曲应力", "65.651", "124.811", "234.300", "0.533"],
        ["事故", "剪切应力", "14.922", "28.370", "219.396", "0.129"],
    ]


def test_appendix_page_break_cleanup_ignores_toc_entries():
    document = Document()
    if "TOC 1" not in [style.name for style in document.styles]:
        document.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
    toc = document.add_paragraph("附录C：焊缝评定原理\t26")
    toc.style = "TOC 1"
    front_break = document.add_paragraph()
    front_break.add_run().add_break(WD_BREAK.PAGE)
    document.add_paragraph("概述")
    document.add_paragraph("附录A：模态分析结果")
    appendix_break = document.add_paragraph()
    appendix_break.add_run().add_break(WD_BREAK.PAGE)
    document.add_paragraph("图A-1 一阶模态振型图")

    audit = _remove_empty_page_breaks_inside_appendices(document)

    assert audit["removed"] == 1
    assert _paragraph_has_page_break(front_break)
    assert appendix_break._element.getparent() is None


def test_static_report_spectrum_uses_selected_static_elevation(monkeypatch, tmp_path):
    workbook = tmp_path / "spectrum.xlsm"
    workbook.write_text("placeholder", encoding="utf-8")
    input_payload = {
        "project": {"project_code": "1818", "building": "NS环形区", "elevation": 7.5},
        "metadata": {
            "analysis_method": "static",
            "static_elevation_candidates": [7.5],
            "static_acceleration_source": {
                "workbook": str(workbook),
                "selected_elevation": 13.09,
                "elevations": [13.09],
            },
        },
    }

    import core.spectra.response_spectrum_writer as response_spectrum_writer
    import core.spectra.static_coefficients as static_coefficients
    import core.spectra.workbook_envelope as workbook_envelope

    seen_elevations: list[float] = []

    def fake_curve_at_elevation(curves, *, level, direction, elevation, damping):
        seen_elevations.append(float(elevation))
        return SimpleNamespace(
            frequency_hz=(0.1, 1.0),
            acceleration_g=(0.01, 0.02),
            source_ref=f"{level}-{direction}-{elevation}",
        )

    monkeypatch.setattr(static_coefficients, "resolve_segmented_spectrum_sheet", lambda *args, **kwargs: "NS环形区_1818")
    monkeypatch.setattr(static_coefficients, "_read_segmented_sheet", lambda *args, **kwargs: [])
    monkeypatch.setattr(static_coefficients, "_curve_at_elevation", fake_curve_at_elevation)
    monkeypatch.setattr(workbook_envelope, "_envelope_curves", lambda curves, source_ref: curves[0])
    monkeypatch.setattr(
        response_spectrum_writer,
        "_curve_payload",
        lambda curve: [
            {"frequency_hz": frequency, "acceleration_g": acceleration}
            for frequency, acceleration in zip(curve.frequency_hz, curve.acceleration_g)
        ],
    )

    assert _static_report_spectrum_elevations(input_payload) == [13.09]
    curves = _points_from_spectrum_file(tmp_path, input_payload)

    assert curves is not None
    assert {key: len(value) for key, value in curves.items()} == {
        "SL-1_vertical": 2,
        "SL-1_horizontal": 2,
        "SL-2_vertical": 2,
        "SL-2_horizontal": 2,
    }
    assert seen_elevations == [13.09, 13.09, 13.09, 13.09, 13.09, 13.09]
